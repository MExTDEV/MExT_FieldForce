from __future__ import annotations

from datetime import date
from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "manuals" / "assets"
OUTPUT_DIR = ROOT / "output" / "docx"
OUTPUT_FILE = OUTPUT_DIR / "MExT_FieldForce_Complete_Handleiding.docx"

NAVY = "0B315B"
BLUE = "0D4E91"
MID_BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FC"
GREEN = "167A5A"
PALE_GREEN = "EAF7F1"
AMBER = "A56600"
PALE_AMBER = "FFF5DE"
RED = "A12635"
PALE_RED = "FCECEF"
INK = "172033"
MUTED = "5E6B7A"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80,
                     start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def create_abstract(abstract_id: int, fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def create_num(num_id: int, abstract_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create_abstract(next_abstract, "bullet", "•", "Symbol")
    create_num(next_num, next_abstract)
    create_abstract(next_abstract + 1, "decimal", "%1.")
    create_num(next_num + 1, next_abstract + 1)
    return next_num, next_num + 1


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


class ManualBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self.bullet_num_id, self.decimal_num_id = add_numbering_definitions(self.doc)
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.font.size = Pt(11)
        normal.font.color.rgb = rgb(INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        for name, size, color, before, after in (
            ("Heading 1", 16, MID_BLUE, 18, 10),
            ("Heading 2", 13, MID_BLUE, 14, 7),
            ("Heading 3", 12, DARK_BLUE, 10, 5),
        ):
            style = styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = rgb(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        caption = styles["Caption"]
        caption.font.name = "Calibri"
        caption.font.size = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = rgb(MUTED)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.keep_with_next = False

        if "Code" not in styles:
            code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        else:
            code = styles["Code"]
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(8.5)
        code.font.color.rgb = rgb(INK)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Inches(0.2)

        self.doc.core_properties.title = "MExT FieldForce - Complete handleiding"
        self.doc.core_properties.subject = "Gebruikershandleiding en technisch framework voor Coaching, SalesDay en Contract"
        self.doc.core_properties.author = "MExT FieldForce"
        self.doc.core_properties.keywords = "FieldForce, Coaching, SalesDay, Contract, handleiding, architectuur"

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("MExT FieldForce  |  Complete handleiding")
        set_run_font(r, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        add_page_field(footer.paragraphs[0])

        settings = self.doc.settings._element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    def title_page(self) -> None:
        logo = ROOT / "public" / "assets" / "fieldforce-logo.png"
        if logo.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            shape = run.add_picture(str(logo), width=Inches(2.7))
            shape._inline.docPr.set("descr", "Logo van MExT FieldForce")
            p.paragraph_format.space_after = Pt(72)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run("PRODUCT- EN PROCESHANDLEIDING")
        set_run_font(r, size=10, color=AMBER, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("MExT FieldForce")
        set_run_font(r, size=30, color=NAVY, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("Coaching · SalesDay · Contract")
        set_run_font(r, size=17, color=BLUE, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(56)
        r = p.add_run("Gebruikershandleiding, procesboek en technisch framework")
        set_run_font(r, size=12, color=MUTED, italic=True)

        meta = self.doc.add_table(rows=4, cols=2)
        set_table_geometry(meta, [2700, 6660])
        data = [
            ("Versie", "1.0"),
            ("Documentdatum", "17 juli 2026"),
            ("Doelgroep", "Vertegenwoordigers, verkoopleiders, management, backoffice, beheerders en ontwikkelaars"),
            ("Status", "Interne handleiding op basis van de broncode en goedgekeurde modulebeslissingen"),
        ]
        for row, (label, value) in zip(meta.rows, data):
            set_cell_shading(row.cells[0], LIGHT_BLUE)
            rp = row.cells[0].paragraphs[0]
            rr = rp.add_run(label)
            set_run_font(rr, size=9.5, color=NAVY, bold=True)
            vp = row.cells[1].paragraphs[0]
            vr = vp.add_run(value)
            set_run_font(vr, size=9.5, color=INK)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(34)
        r = p.add_run("INTERN GEBRUIK · screenshots kunnen testdata en echte gebruikersnamen bevatten")
        set_run_font(r, size=8.5, color=RED, bold=True)
        self.doc.add_page_break()

    def h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def p(self, text: str, bold_prefix: str | None = None) -> None:
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            r1 = p.add_run(bold_prefix)
            set_run_font(r1, bold=True, color=INK)
            r2 = p.add_run(text[len(bold_prefix):])
            set_run_font(r2, color=INK)
        else:
            r = p.add_run(text)
            set_run_font(r, color=INK)

    def bullet(self, text: str) -> None:
        p = self.doc.add_paragraph()
        set_num(p, self.bullet_num_id)
        r = p.add_run(text)
        set_run_font(r, color=INK)

    def numbered(self, text: str) -> None:
        p = self.doc.add_paragraph()
        set_num(p, self.decimal_num_id)
        r = p.add_run(text)
        set_run_font(r, color=INK)

    def callout(self, title: str, text: str, tone: str = "info") -> None:
        colors = {
            "info": (PALE_BLUE, BLUE),
            "success": (PALE_GREEN, GREEN),
            "warning": (PALE_AMBER, AMBER),
            "risk": (PALE_RED, RED),
        }
        fill, accent = colors[tone]
        table = self.doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [TABLE_WIDTH_DXA])
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        r = p.add_run(title + "\n")
        set_run_font(r, size=10.5, color=accent, bold=True)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def table(self, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        set_table_geometry(table, widths)
        set_repeat_table_header(table.rows[0])
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(header)
            set_run_font(r, size=9, color=NAVY, bold=True)
        for row_data in rows:
            row = table.add_row()
            for index, value in enumerate(row_data):
                p = row.cells[index].paragraphs[0]
                r = p.add_run(str(value))
                set_run_font(r, size=9, color=INK)
        set_table_geometry(table, widths)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def screenshot(self, filename: str, caption: str, alt: str,
                   width: float = 6.25) -> None:
        path = ASSET_DIR / filename
        if not path.exists():
            self.callout("Screenshot niet beschikbaar", f"Het bestand {filename} kon niet worden opgenomen.", "warning")
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        shape = run.add_picture(str(path), width=Inches(width))
        shape._inline.docPr.set("descr", alt)
        cp = self.doc.add_paragraph(style="Caption")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_run_font(cr, size=9, color=MUTED, italic=True)

    def code(self, text: str) -> None:
        p = self.doc.add_paragraph(style="Code")
        r = p.add_run(text)
        set_run_font(r, name="Consolas", size=8.5, color=INK)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def toc(self) -> None:
        """Insert a Word TOC field that Word/LibreOffice can update on render/open."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Inhoudsopgave wordt bij openen of renderen bijgewerkt."
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, placeholder, end])

    def save(self) -> None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        self.doc.save(OUTPUT_FILE)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int],
                title: str, body: str, fill: str, outline: str,
                title_color: str = NAVY) -> None:
    draw.rounded_rectangle(box, radius=24, fill="#" + fill, outline="#" + outline, width=3)
    x1, y1, x2, y2 = box
    draw.text((x1 + 24, y1 + 20), title, font=font(27, True), fill="#" + title_color)
    wrapped = textwrap.fill(body, width=max(22, int((x2 - x1) / 15)))
    draw.multiline_text((x1 + 24, y1 + 64), wrapped, font=font(19), fill="#" + INK, spacing=7)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
          color: str = BLUE, width: int = 6) -> None:
    draw.line([start, end], fill="#" + color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    left = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    right = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill="#" + color)


def diagram_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(38, True), fill="#" + NAVY)
    draw.text((70, 100), subtitle, font=font(21), fill="#" + MUTED)
    return image, draw


def save_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    image, draw = diagram_canvas("Platformarchitectuur", "Eén shell, gedeelde beveiliging en module-eigen workflows")
    rounded_box(draw, (70, 180, 1530, 300), "FieldForce platform", "Aanmelding · gebruikers · rollen · rechten · landen/teams · taal · navigatie · audit", PALE_BLUE, MID_BLUE)
    boxes = [
        (90, 380, 500, 570, "Coaching", "Planning, begeleiding, actiepunten, contactmomenten en hulp"),
        (595, 380, 1005, 570, "SalesDay + Inventory", "Daguitvoering, klant, verkoop, voorraad, cash en offline sync"),
        (1100, 380, 1510, 570, "Contract", "Calculatie, catalogus, snapshots, ondertekening en brief"),
    ]
    for x1, y1, x2, y2, title, body in boxes:
        rounded_box(draw, (x1, y1, x2, y2), title, body, WHITE, LIGHT_BLUE)
        arrow(draw, ((x1 + x2) // 2, 300), ((x1 + x2) // 2, 375))
    rounded_box(draw, (190, 690, 650, 830), "Prisma + MariaDB", "Transacties, snapshots, audit en scopebewuste data", PALE_GREEN, GREEN)
    rounded_box(draw, (950, 690, 1410, 830), "Externe diensten", "Entra ID · Outlook · BC/NAV · toekomstig Odoo · Power BI", PALE_AMBER, AMBER)
    arrow(draw, (800, 570), (430, 685))
    arrow(draw, (800, 570), (1180, 685))
    image.save(ASSET_DIR / "diagram-platform-architectuur.png")

    image, draw = diagram_canvas("Begeleiding: levenscyclus", "Eén record, één formulier en één goedkeuringspad")
    stages = [
        ("Gepland", "Datum, persoon, tijd, melding en focus"),
        ("In uitvoering", "Voorbereiden, observeren, scoren en afspraken"),
        ("Onvolledig", "Veilig bewaren en later hervatten"),
        ("Wachten op akkoord", "Reflectie door gecoachte persoon; dossier vergrendeld"),
        ("Afgerond", "Officiële historiek, rapport en vergelijking"),
    ]
    x = 45
    for index, (title, body) in enumerate(stages):
        fill = PALE_GREEN if index == len(stages) - 1 else PALE_BLUE
        outline = GREEN if index == len(stages) - 1 else MID_BLUE
        rounded_box(draw, (x, 300, x + 270, 560), title, body, fill, outline)
        if index < len(stages) - 1:
            arrow(draw, (x + 270, 430), (x + 320, 430))
        x += 320
    draw.text((405, 690), "Bij een correctie: aanvraag intrekken → aanpassen → opnieuw ter akkoord aanbieden", font=font(23, True), fill="#" + RED)
    image.save(ASSET_DIR / "diagram-coaching-lifecycle.png")

    image, draw = diagram_canvas("SalesDay: werkdag", "Van synchronisatie en voorbereiding tot dagafsluiting")
    stages = [
        ("1. Toegang", "Actief toestel, online sessie, geen oude sync, kascontrole"),
        ("2. Voorbereiding", "Volgende werkdag vanaf landparameter; notities en aanbevelingen"),
        ("3. Agenda", "Verplichte ERP-volgorde; klantdossier per afspraak"),
        ("4. Uitvoering", "Klantgegevens, bezoek, verkoop, lead, opvolging en referentie"),
        ("5. Afsluiten", "Elke afspraak heeft uitkomst; dag sluiten en resterende sync tonen"),
    ]
    y = 190
    for index, (title, body) in enumerate(stages):
        x1 = 90 if index % 2 == 0 else 840
        x2 = 760 if index % 2 == 0 else 1510
        rounded_box(draw, (x1, y, x2, y + 120), title, body, WHITE, MID_BLUE)
        if index < len(stages) - 1:
            next_x = 840 if (index + 1) % 2 == 1 else 90
            arrow(draw, ((x1 + x2) // 2, y + 120), ((next_x + (1510 if next_x == 840 else 760)) // 2, y + 165))
        y += 145
    image.save(ASSET_DIR / "diagram-salesday-workday.png")

    image, draw = diagram_canvas("SalesDay synchronisatie", "Duurzame, idempotente uitwisseling zonder stille dataverliezen")
    rounded_box(draw, (60, 250, 440, 570), "Tablet / PWA", "Versleutelde replica, drafts en wachtrij. Automatisch syncen; handmatig opnieuw proberen.", PALE_BLUE, MID_BLUE)
    rounded_box(draw, (610, 210, 990, 610), "FieldForce server", "API + scopecontrole\n\nOutbox voor opdrachten\nInbox voor ERP-events\nReconciliatie en incidenten", WHITE, NAVY)
    rounded_box(draw, (1160, 250, 1540, 570), "ERP", "Business Central/NAV nu; Odoo later. Finale effectieve bron na bevestiging.", PALE_AMBER, AMBER)
    arrow(draw, (440, 350), (600, 350), GREEN)
    arrow(draw, (600, 480), (440, 480), BLUE)
    arrow(draw, (990, 350), (1150, 350), GREEN)
    arrow(draw, (1150, 480), (990, 480), BLUE)
    draw.text((275, 650), "Lokaal commando + businessmutatie in één transactie", font=font(21, True), fill="#" + GREEN)
    draw.text((940, 710), "Timeout is nooit automatisch succes: status opvragen of reconciliëren", font=font(21, True), fill="#" + RED)
    image.save(ASSET_DIR / "diagram-salesday-sync.png")

    image, draw = diagram_canvas("Contract: van klant naar ondertekend document", "Vier stappen met serverberekening en onveranderlijke snapshots")
    stages = [
        ("Klant", "Bestaande relatie kiezen of klant vastleggen; taal bepalen"),
        ("Artikelen", "Catalogus zoeken, aantallen ingeven en regels controleren"),
        ("Looptijd", "3 of 5 jaar; actieve modelregel bepaalt de korting"),
        ("Afronden", "Naam geven, concept opslaan, interne PDF of ondertekenen"),
        ("Bewijs", "Actieve taaltemplate, handtekening, immutable PDF en versie"),
    ]
    x = 45
    for index, (title, body) in enumerate(stages):
        fill = PALE_GREEN if index == 4 else WHITE
        outline = GREEN if index == 4 else MID_BLUE
        rounded_box(draw, (x, 290, x + 270, 570), title, body, fill, outline)
        if index < len(stages) - 1:
            arrow(draw, (x + 270, 430), (x + 320, 430))
        x += 320
    draw.text((305, 690), "De server is de rekenkundige bron; prijzen, kosten, klanttaal en modelversie worden gesnapshot.", font=font(22, True), fill="#" + NAVY)
    image.save(ASSET_DIR / "diagram-contract-flow.png")

    image, draw = diagram_canvas("Effectieve toegang", "Een rolnaam alleen is nooit voldoende")
    stages = [
        ("Rol actief?", "Systeemrol en runtimeconfiguratie"),
        ("Module actief?", "Domein en functionele module"),
        ("Recht?", "Roldefault plus gebruikersoverride"),
        ("Scope?", "Land, team, eigen gebruiker en record"),
        ("Status?", "Levenscyclus, dagpoort, toestel en featureflag"),
        ("Toestaan", "Menu, pagina, API en bestand hanteren dezelfde uitkomst"),
    ]
    x = 45
    for index, (title, body) in enumerate(stages):
        fill = PALE_GREEN if index == 5 else WHITE
        outline = GREEN if index == 5 else MID_BLUE
        rounded_box(draw, (x, 300, x + 225, 560), title, body, fill, outline)
        if index < len(stages) - 1:
            arrow(draw, (x + 225, 430), (x + 265, 430))
        x += 265
    image.save(ASSET_DIR / "diagram-effective-access.png")


def add_intro(builder: ManualBuilder) -> None:
    b = builder
    b.h1("1. Leeswijzer, status en gebruik")
    b.p("Deze handleiding beschrijft MExT FieldForce als één platform met drie hoofdonderdelen: Coaching, SalesDay en Contract. Ze combineert dagelijkse werkwijzen met procesregels, rechten, gegevensstromen, beheer en technische architectuur. De beschrijving volgt de broncode en de goedgekeurde modulebeslissingen zoals ze op 17 juli 2026 in de repository aanwezig waren.")
    b.callout("Belangrijk over de huidige omgeving", "Coaching bevat hoofdzakelijk testdata; de gebruikers zijn echte accounts. SalesDay is in broncode geïmplementeerd maar voor de vastgelegde lokale scope niet geactiveerd en niet gevuld met UAT-mockdata. De concrete Business Central/NAV-koppeling is nog extern te realiseren. Contract is native aanwezig, maar de live schermafbeelding toont de server-side laadstatus. Er is voor deze handleiding geen database, featureflag, productieomgeving of ERP gewijzigd.", "warning")
    b.h2("1.1 Doelgroepen")
    for item in [
        "Vertegenwoordigers: dagelijkse uitvoering, eigen afspraken, eigen dossiers, acties, verkoop, voorraad en cash.",
        "Verkoopleiders en management: plannen, begeleiden, opvolgen en lezen binnen effectieve team- of landenscope.",
        "Backoffice en magazijn: ERP-bevestigingen, voorraad- en kasprocessen, reconciliatie en operationele ondersteuning.",
        "Administrators en Super Admins: gebruikers, rollen, modules, parameters, integraties en technische incidenten.",
        "Ontwikkelaars en functioneel beheerders: architectuur, gegevensbezit, teststrategie, migraties en wijzigingsimpact.",
    ]:
        b.bullet(item)
    b.h2("1.2 Statuslegenda")
    b.table(
        ["Status", "Betekenis voor gebruik en documentatie"],
        [
            ["DEFINED", "Het zakelijke gedrag is voldoende vastgelegd. De handleiding beschrijft het proces als normatief."],
            ["PARTIALLY_DEFINED", "Alleen het beschreven deel mag als vast proces worden gebruikt; resterende stappen vragen een zakelijke beslissing."],
            ["UNDEFINED", "Geen workflow, status, formule of recht afleiden uit een menu of technisch model."],
            ["IMPLEMENTED IN SOURCE", "Broncode en tests bestaan, maar database-uitrol, integratie of externe acceptatie kan nog ontbreken."],
            ["EXTERNAL ACCEPTANCE PENDING", "De softwarebasis bestaat; productie vereist nog bewijs buiten de repository."],
        ],
        [2200, 7160],
    )
    b.h2("1.3 Wat deze handleiding wel en niet belooft")
    b.p("De handleiding legt bestaand en goedgekeurd gedrag uit. Ze maakt geen onbekende ERP-endpoints, printerdrivers, rapporteringsformules, retentieperioden of niet-goedgekeurde Coaching-workflows aan. Waar functionaliteit nog niet actief of zakelijk gedefinieerd is, staat dat expliciet aangegeven.")
    b.screenshot(
        "00-applicatiewisselaar.png",
        "Figuur 1 — De applicatiewisselaar bundelt actieve modules en directe links. De getoonde gegevens komen uit de lokale testomgeving.",
        "Applicatiewisselaar van MExT FieldForce met Coaching, Inventory, PST, Contract en Service.",
    )
    b.h2("1.4 Snelle start")
    for item in [
        "Meld aan met het toegekende FieldForce-account; Microsoft Entra ID is het voorkeursmechanisme.",
        "Controleer rechtsboven taal, naam, rol en land. Kies NL, FR of DE wanneer nodig.",
        "Klik op de gebruikersnaam om van applicatie te wisselen of een directe link te openen.",
        "Gebruik de donkerblauwe linkerzijbalk voor de actieve module. Alleen toegestane items worden getoond.",
        "Behandel een ontbrekend menu niet automatisch als een fout: moduleactivatie, rolrecht, gebruikersoverride, scope en workflowstatus bepalen samen de zichtbaarheid.",
    ]:
        b.numbered(item)


def add_platform(builder: ManualBuilder) -> None:
    b = builder
    b.h1("2. Platform, navigatie en effectieve toegang")
    b.screenshot("diagram-platform-architectuur.png", "Figuur 2 — Functionele platformarchitectuur.", "Diagram van de FieldForce platformarchitectuur.")
    b.h2("2.1 Eén platform, module-eigen processen")
    b.p("FieldForce gebruikt één login, één gebruikers- en teammodel, één rechtenmodel, één navigatieshell, één vertaalfundament, één auditbasis en één Prisma/MariaDB-database. Coaching, SalesDay en Contract blijven wel eigenaar van hun eigen businessobjecten en levenscycli. Een Planning-scherm mag een Begeleiding tonen, maar Planning mag nooit een tweede Begeleiding of afwijkend statuspad creëren.")
    b.h2("2.2 Navigatie")
    for item in [
        "Linkerzijbalk: snelle navigatie binnen het actieve domein.",
        "Applicatiewisselaar: openen via de gebruikersnaam; toont actieve domeinen en directe links.",
        "Taalkeuze: NL, FR of DE; nieuwe zichtbare teksten moeten via het vertaalsysteem lopen.",
        "Meldingen: de bel toont alleen records die de gebruiker effectief mag openen.",
        "Diepe links: een URL is geen omweg rond rechten. Pagina en API controleren dezelfde toegang opnieuw.",
    ]:
        b.bullet(item)
    b.h2("2.3 Effectieve toegang")
    b.screenshot("diagram-effective-access.png", "Figuur 3 — Toegang wordt laag per laag berekend.", "Diagram van de effectieve toegangscontrole.")
    b.table(
        ["Laag", "Voorbeelden", "Waarom ze afzonderlijk telt"],
        [
            ["Rolconfiguratie", "Representative, Sales Leader, Sales Manager, Country Manager, Admin", "Bepaalt defaults, niet automatisch alle rechten."],
            ["Moduleactivatie", "Coaching, SalesDay, Inventory, Contract", "Een gedeactiveerde module blijft onbereikbaar, zelfs met een menu-recht."],
            ["Gebruikersoverride", "Expliciet toestaan of weigeren", "Een individuele afwijking kan de rol-default wijzigen."],
            ["Organisatiescope", "Eigen gebruiker, team, land(en), globaal", "Beperkt welke records zichtbaar of muteerbaar zijn."],
            ["Levenscyclus", "Gepland, wachten op akkoord, afgerond", "Een zichtbaar record kan toch vergrendeld zijn."],
            ["Operationele poort", "Toestel, sync dag -1, cash, featureflags", "SalesDay gebruikt bijkomende dag- en toestelsleutels."],
        ],
        [1750, 3000, 4610],
    )
    b.h2("2.4 Rollen in één oogopslag")
    b.table(
        ["Rol", "Primaire scope", "Typisch gedrag"],
        [
            ["Vertegenwoordiger", "Eigen gegevens", "Eigen coaching lezen/goedkeuren, eigen acties en volledige toegestane SalesDay-werkdag."],
            ["Verkoopleider", "Eigen team", "Coaching plannen/uitvoeren, team opvolgen; SalesDay operationeel alleen-lezen."],
            ["Sales Manager", "Toegewezen landen", "Coacht Representatives en Verkoopleiders; managementinzichten binnen scope."],
            ["Country Manager", "Toegewezen land(en)", "Landmanagement, coaching en opvolging volgens rechten."],
            ["Group Manager", "Expliciete scope", "Gedeeltelijk gedefinieerd; nooit automatisch gelijk aan Super Admin."],
            ["Admin", "Toegewezen land(en)", "Gebruikers/configuratie en zichtbare ontgrendelde records; geen onbeperkte impersonatie."],
            ["Super Admin", "Globaal", "Volledige zichtbaarheid en beheerrechten, maar nog steeds gebonden aan zakelijke locks."],
            ["Service Operator", "Expliciete scope", "Gedeeltelijk gedefinieerd binnen Coaching."],
        ],
        [1700, 2050, 5610],
    )
    b.callout("Praktische diagnose", "Ziet iemand een pagina of knop niet? Controleer in deze volgorde: rol actief → module actief → rolrecht → gebruikersoverride → land/team/eigen scope → recordstatus → eventuele SalesDay-feature- of dagpoort.", "info")


def add_coaching(builder: ManualBuilder) -> None:
    b = builder
    b.h1("3. Coaching — complete gebruikershandleiding")
    b.p("Coaching ondersteunt planning, uitvoering, documentatie, reflectie, goedkeuring en opvolging van medewerkers in het veld. Het centrale uitgangspunt is één betrouwbaar dossier per interventie, ongeacht of de gebruiker het opent via Dashboard, Planning, Begeleidingen, Mijn Team of een melding.")
    b.h2("3.1 Functionele kaart")
    b.table(
        ["Onderdeel", "Status", "Kernfunctie"],
        [
            ["Dashboard", "DEFINED", "Vandaag, aandacht, prioriteiten, teamindicatoren en snelle toegang."],
            ["Mijn Team", "DEFINED", "Scopebewust medewerkersoverzicht en fiche."],
            ["Planning", "DEFINED als presentatie", "Agendaweergave; opent het object dat de workflow bezit."],
            ["Begeleidingen", "DEFINED", "Plannen, voorbereiden, uitvoeren, reflecteren, goedkeuren en historiek."],
            ["Actiepunten", "PARTIALLY_DEFINED", "Definities en operationeel afsluiten; heropenen/herverdelen niet vastgelegd."],
            ["Contactmomenten", "DEFINED", "Verslag zonder volledig scoreformulier; delen vergrendelt."],
            ["Hulpaanvragen", "DEFINED", "Ondersteuningsvraag met concrete respons of opvolging; geen chat."],
            ["Tussentijdse evaluaties", "PARTIALLY_DEFINED", "Startersevaluaties en manuele start; verdere akkoordflow nog open."],
            ["Retrainingen", "UNDEFINED", "Geen proces verzinnen."],
            ["Salestrainingen", "UNDEFINED", "Geen proces verzinnen."],
            ["Rapportage", "UNDEFINED", "Geen formules, rankings of drempels verzinnen."],
        ],
        [2350, 1750, 5260],
    )
    b.h2("3.2 Dashboard")
    b.p("Het Dashboard is de operationele startpagina. Het toont wat vandaag moet gebeuren, wat al uitgevoerd is, hoeveel begeleidingen of actiepunten aandacht vragen en — voor managementgebruikers — prioriteiten en teamindicatoren binnen scope. Elke kaart linkt naar de eigenaar van het proces; het Dashboard bewaart geen tweede kopie van een workflow.")
    b.screenshot("01-coaching-dashboard.png", "Figuur 4 — Coaching Dashboard met aandachtspunten, prioriteiten en kerncijfers.", "Coaching Dashboard van FieldForce.")
    for item in [
        "Open een item onder Uit te voeren om het bestaande dossier te starten of voort te zetten.",
        "Gebruik Nieuwe begeleiding alleen wanneer u voor de gekozen medewerker planningsrecht en scope hebt.",
        "Klik op een KPI- of aandachtstegel om de onderliggende lijst te openen; aantallen zijn geen afzonderlijke rapportbron.",
        "Een lege kaart betekent niet automatisch een laadfout. Controleer de tekst en de ingestelde filters/scope.",
    ]:
        b.bullet(item)
    b.h2("3.3 Mijn Team")
    b.p("Mijn Team is niet zichtbaar voor vertegenwoordigers. Een Verkoopleider ziet het eigen team; hogere rollen zien alleen medewerkers binnen de toegewezen landen of expliciete scope. Teams zonder verkoopleider blijven zichtbaar voor bevoegd management, maar worden niet automatisch aan andere verkoopleiders toegewezen.")
    b.h3("Medewerkersfiche")
    b.p("De fiche kan Overzicht, Prestatiecirkel, Persoonlijke Criteria, Actiepunten, Hulpaanvragen, Begeleidingen, KPI's, Evaluaties, Contactmomenten, Retrainingen, Sales Trainingen en Tijdlijn tonen. Alleen actieve modules en toegestane secties verschijnen. E-mail en telefoon staan in de kop als semantische links. Een ontbrekende score is een neutraal streepje, nooit automatisch nul.")
    b.h2("3.4 Planning")
    b.p("Planning is een kalenderpresentatie voor FieldForce-items en, waar gekoppeld, Outlook-items. Het scherm maakt geen eigen kopieën: een Begeleiding opent het Begeleidingsdossier, een Contactmoment opent het Contactmoment en een opvolging vanuit een Hulpaanvraag opent het aangemaakte doelobject.")
    b.screenshot("04-coaching-planning.png", "Figuur 5 — Weekweergave in Planning; itemtypen krijgen een eigen visuele aanduiding.", "Weekplanning in FieldForce.")
    for item in [
        "Schakel tussen dag, week en maand volgens de gewenste tijdschaal.",
        "De hoogte en plaats van een item volgen de werkelijke duur.",
        "Een verrassingsbegeleiding mag niet via Planning uitlekken naar de gecoachte persoon.",
        "Outlook-synchronisatie is ondersteunend. De FieldForce-record blijft de zakelijke bron.",
    ]:
        b.bullet(item)
    b.h2("3.5 Begeleiding plannen")
    b.screenshot("diagram-coaching-lifecycle.png", "Figuur 6 — Levenscyclus van een Begeleiding.", "Diagram van de levenscyclus van een Begeleiding.")
    b.h3("Voorwaarden")
    for item in [
        "De planner is aangemeld en heeft het recht om een Begeleiding te creëren.",
        "De gecoachte persoon valt binnen de effectieve scope.",
        "Datum, begin- en eindtijd zijn geldig.",
        "De vereiste focusgebieden of criteria zijn gekozen.",
    ]:
        b.bullet(item)
    b.h3("Werkwijze")
    for item in [
        "Open Begeleidingen en kies Nieuwe begeleiding.",
        "Selecteer de gecoachte persoon. Een hogere manager kan ook een Verkoopleider begeleiden wanneer rechten en scope dit toelaten.",
        "Vul datum en tijd in en bepaal of de persoon vooraf wordt ingelicht.",
        "Selecteer de focusgebieden. Die keuze bepaalt welke criteria tijdens uitvoering verschijnen.",
        "Kies in de voorbereiding één volledig afgeronde eerdere Begeleiding als referentie. De nieuwste geldige is standaard geselecteerd.",
        "Controleer de samenvatting en sla op. FieldForce wacht op bevestiging van de create-opdracht voordat het terugnavigeert.",
        "Een Outlook-fout mag het geldige FieldForce-dossier niet terugdraaien; alleen de synchronisatiestatus wijzigt.",
    ]:
        b.numbered(item)
    b.callout("Verrassingsbegeleiding", "Wanneer vooraf melden uit staat, mag de gecoachte vertegenwoordiger het geplande dossier niet zien. Het wordt pas zichtbaar bij de afgesproken goedkeuringsfase. Managementzicht blijft scope- en rechtengedreven.", "warning")
    b.h2("3.6 Begeleidingenoverzicht")
    b.p("Het overzicht groepeert dossiers in Vandaag, Toekomst, Onvolledig, Wachten op akkoord en Uitgevoerd. Groepering volgt beschikbare landen, teams en personen. Filters versmallen de lijst, maar verruimen nooit de server-side scope.")
    b.screenshot("02-coaching-begeleidingen.png", "Figuur 7 — Begeleidingen met status, synchronisatie en dossieractie.", "Overzicht van Begeleidingen in FieldForce.")
    b.h2("3.7 Voorbereiding en historische vergelijking")
    for item in [
        "De geselecteerde referentie moet volledig afgerond, van dezelfde persoon, in het verleden en binnen scope zijn.",
        "De gekozen referentie blijft bewaard wanneer de planning later opnieuw wordt geopend.",
        "Historische labels, criteria, volgorde, scores en opmerkingen komen uit snapshots; huidige configuratie herschrijft het verleden niet.",
        "De nieuwste afgeronde Begeleiding en de gekozen referentie kunnen naast elkaar in de voorbereiding-PDF verschijnen.",
        "Open actiepunten blijven live gegevens; een oudere referentie verandert hun huidige status niet.",
    ]:
        b.bullet(item)
    b.h2("3.8 Uitvoering: het zevenstappenverslag")
    b.screenshot("03-coaching-dossier-stappen.png", "Figuur 8 — Het vrij navigeerbare zevenstappenverslag met autosave-status.", "Begeleidingsdossier met zeven stappen en synchronisatiestatus.")
    b.table(
        ["Stap", "Doel", "Belangrijke controle"],
        [
            ["1. Algemene gegevens", "Gepland en werkelijk bezoekcontext vastleggen.", "Aankomst, vertrek en kilometers blijven zakelijke invoer."],
            ["2. Voorbereiding", "Historiek, referentie, scores en open opvolging raadplegen.", "Historische snapshots zijn read-only."],
            ["3. Evaluatie algemene punten", "Expliciete numerieke score of NVT per verplicht criterium.", "Niet gekozen is anders dan NVT."],
            ["4. Evaluatie persoonlijkheid", "Gedrags- en persoonlijkheidscriteria evalueren.", "Ook hier is expliciete keuze verplicht voor afronding."],
            ["5. Afspraken", "Eén of meer klantenbezoeken en observaties registreren.", "Minstens één bezoek voor completion."],
            ["6. Actiepunten", "Concrete opvolging voor de gecoachte persoon vastleggen.", "Minstens één actiepunt voor een voltooide Begeleiding."],
            ["7. Afsluiten", "Samenvatting en ontbrekende verplichte gegevens controleren.", "Finaliseren stuurt daarna naar akkoord; opslaan/sluiten niet."],
        ],
        [1850, 3750, 3760],
    )
    b.h3("Autosave en herstel")
    for item in [
        "Alle veldtypen gebruiken één seriële opslagwachtrij; tekstinvoer heeft een korte debounce.",
        "Vóór stapnavigatie worden nog niet verzonden wijzigingen geflusht.",
        "Een oudere serverrespons mag een nieuwere lokale versie nooit overschrijven.",
        "Bij een fout blijft de lokale waarde en browserdraft bestaan en verschijnt een niet-blokkerende retry.",
        "Een herladen bewerkbaar rapport herstelt een nieuwere niet-verzonden draft.",
    ]:
        b.bullet(item)
    b.h2("3.9 Ter akkoord aanbieden, reflectie en goedkeuring")
    b.h3("Indienen")
    b.p("Finaliseren valideert alle verplichte gegevens server-side, bewaart het dossier en start de bestaande overgang naar Wachten op akkoord. Vanaf dat moment is het dossier read-only. De gecoachte persoon ontvangt een taak, in-app melding en best-effort e-mail.")
    b.h3("Reflectie door de gecoachte persoon")
    for item in [
        "De gecoachte persoon beantwoordt eerst drie verplichte WYSIWYG-reflectievragen.",
        "Zolang niet alle antwoorden betekenisvolle tekst bevatten, blijven rapportinhoud en akkoordknoppen geblokkeerd.",
        "De antwoorden worden server-side gesaneerd en staan op de bestaande Approval-record.",
        "Na akkoord of niet-akkoord zijn de antwoorden read-only.",
        "Managers met dossierzicht zien de reflectie read-only; ontbrekende antwoorden worden duidelijk gemarkeerd.",
    ]:
        b.bullet(item)
    b.h3("Correctie na indiening")
    b.p("Een bevoegde coach trekt eerst de akkoordaanvraag in. Het dossier keert terug naar een bewerkbare status, behoudt auditinformatie en moet na wijziging opnieuw worden ingediend. Een afgerond dossier wordt nooit stilzwijgend ontgrendeld.")
    b.h2("3.10 Actiepunten")
    b.p("Actiepunten kunnen als globale, land-, team- of gebruikersdefinitie bestaan. Concrete punten ontstaan onder meer uit een Begeleiding of Contactmoment. Sluiten raakt alleen de concrete gebruikerstoewijzing en nooit de gedeelde definitie, de oorspronkelijke score, het ondertekende verslag of de historische PDF.")
    b.screenshot("05-coaching-actiepunten.png", "Figuur 9 — Gegroepeerde Actiepunten met gebruiker, scope, status en prioriteit.", "Actiepuntenoverzicht in FieldForce.")
    b.h3("Sluiten")
    for item in [
        "Vereist het functionele recht actionPointsClose en de juiste module-/menuscope.",
        "Verkoopleider: alleen concrete punten voor de eigen toegestane teamleden.",
        "Country Manager en Admin: alleen toegewezen landen.",
        "Group Manager: alleen expliciet toegewezen scope; niet automatisch globaal.",
        "Super Admin: globaal, maar de actie blijft auditeerbaar.",
        "Dubbel sluiten is idempotent; bestaande sluitmetadata blijft behouden.",
        "Heropenen, goedkeuren en herverdelen zijn nog niet zakelijk gedefinieerd.",
    ]:
        b.bullet(item)
    b.h2("3.11 Contactmomenten")
    b.p("Een Contactmoment is een geplande managementinteractie zonder het volledige Coaching-scoreformulier. De planner kiest persoon, datum, tijd, eventuele voorafmelding en context. Voor delen wordt een WYSIWYG-verslag opgesteld; foto's en gebruikersactiepunten zijn optioneel.")
    b.h3("Fotoregels")
    for item in [
        "JPEG, PNG en WebP; maximaal 8 MB per bestand en maximaal 20 foto's per Contactmoment.",
        "MIME-type én bestandssignatuur worden gecontroleerd.",
        "Foto's worden privé opgeslagen en alleen via de geauthenticeerde API gelezen.",
        "Na delen zijn verslag en foto's onveranderlijk voor gewone bewerking.",
        "De PDF neemt de foto's in opgeslagen volgorde op en toont een gecontroleerde placeholder bij een ontbrekend bestand.",
    ]:
        b.bullet(item)
    b.h3("Delen")
    b.p("Delen geeft de doelpersoon toegang, maakt geen goedkeuringstaak, legt tijdstip en actor vast en vergrendelt verslag en galerij. Een toekomstige correctie vraagt een afzonderlijk goedgekeurd proces.")
    b.h2("3.12 Hulpaanvragen")
    b.p("Een vertegenwoordiger vraagt ondersteuning aan de verantwoordelijke Verkoopleider. De aanvraag is geen chat en mag niet met een algemene weigering worden afgesloten. De manager geeft een concrete respons, plant een geschikte opvolging of documenteert een andere actie.")
    b.table(
        ["Fase", "Actie", "Regel"],
        [
            ["Nieuw", "Vertegenwoordiger dient onderwerp en gesaneerde rijke beschrijving in.", "Verantwoordelijke wordt server-side afgeleid; melding en best-effort e-mail."],
            ["In behandeling", "Manager beoordeelt en kiest uit opvolging, concrete respons of één gecontroleerde responsvraag.", "De oorspronkelijke aanvraag wordt niet gewijzigd."],
            ["Respons", "Aanvrager mag één gecontroleerd antwoord geven.", "Dit is een workflowbeurt, geen onbeperkte conversatie."],
            ["Opvolging gepland", "Begeleiding of Contactmoment wordt met de aanvrager vooringevuld.", "Het doelproces voert eigen rechten en validatie opnieuw uit."],
            ["Gesloten", "Concrete uitkomst en actor/tijd worden bewaard.", "Geen functioneel lege afsluittekst; geen status Afgewezen."],
        ],
        [1500, 4200, 3660],
    )
    b.screenshot("06-coaching-hulpaanvragen.png", "Figuur 10 — Een gedeactiveerde functionele module meldt dit expliciet; activatie gebeurt via technisch modulebeheer.", "FieldForce melding dat Hulpaanvragen niet actief is.")
    b.h2("3.13 Tussentijdse evaluaties")
    for item in [
        "Automatische startersevaluaties gelden voor actieve Representatives met niveau STARTER en een officiële starterStartDate.",
        "Vaste momenten: zes kalenderweken, drie kalendermaanden en vijf kalendermaanden na de startdatum.",
        "Manuele evaluaties kunnen voor elke actieve evalueerbare Representative binnen scope worden gestart; het niveau is informatief.",
        "Vragen ondersteunen cumulatieve Global-, Country-, Team- en User-scope en worden bij aanmaak gesnapshot.",
        "Statussen bestaan, maar akkoordflow, definitieve PDF, Outlook-sync en activatie van echte actiepunten zijn nog niet volledig vrijgegeven.",
    ]:
        b.bullet(item)
    b.h2("3.14 Niet-ingevulde domeinen")
    b.callout("Niet als bestaand proces behandelen", "Retrainingen, Salestrainingen en Coaching-rapportage zijn nog UNDEFINED. Een menu, enum of databaseveld is geen goedgekeurde workflow. Ook drempels, rankings, rapportformules, deelnemersregels en aanwezigheidsprocessen mogen niet uit aannames worden ingevuld.", "risk")


def add_salesday(builder: ManualBuilder) -> None:
    b = builder
    b.h1("4. SalesDay — complete proceshandleiding")
    b.p("SalesDay ondersteunt de volledige commerciële werkdag van een vertegenwoordiger: voorbereiding, agenda, klantonderhoud, afspraakuitvoering, verkoopdocumenten, voorraad, cash, dagafsluiting en betrouwbare offline synchronisatie. FieldForce is de invoer-, offline- en presentatielaag; Business Central/NAV en later Odoo worden de finale effectieve bron na bevestiging.")
    b.callout("Huidige productiestatus", "Milestones 1 tot en met 7 zijn in broncode uitgewerkt. De echte BC/NAV-transportlaag, productie-migraties, UAT-sign-off, MDM-oefening en andere externe bewijzen ontbreken nog. De huidige lokale database is bewust niet met mockdata gevuld omdat de naam niet als UAT/testomgeving is gemarkeerd.", "warning")
    b.screenshot("10-salesday-overzicht.png", "Figuur 11 — De featuregate voorkomt toegang wanneer SalesDay niet voor de huidige scope is geactiveerd.", "FieldForce melding dat SalesDay niet geactiveerd is.")
    b.h2("4.1 Dagproces in één beeld")
    b.screenshot("diagram-salesday-workday.png", "Figuur 12 — Goedgekeurde SalesDay-werkdag.", "Diagram van de SalesDay-werkdag.")
    b.h2("4.2 Voorwaarden om de dag te starten")
    for item in [
        "SalesDay, Inventory en benodigde write-flags zijn server-side geactiveerd voor land, team of gebruiker.",
        "De vertegenwoordiger gebruikt het persoonlijk geregistreerde, actieve en correct geprovisioneerde toestel.",
        "De online sessie is geldig; bij hervatten na slaap/lock wordt lokale biometrie of pincode gebruikt.",
        "Er staan geen onbevestigde opdrachten uit dag -1 open. Zo ja, blijven alleen blokkering, sync en support beschikbaar.",
        "Op de eerste effectieve werkdag van de week is het door ERP/backoffice bevestigde kassaldo exact nul.",
        "De agenda van de nieuwe werkdag is succesvol uit het ERP ontvangen.",
    ]:
        b.bullet(item)
    b.h3("Dag -1 blokkering")
    b.p("FieldForce combineert server-outboxbewijzen en de versleutelde lokale wachtrij. Elke opdracht met een eerdere businessdatum die niet ACCEPTED is, blokkeert de volgende werkdag. De gebruiker kiest nooit zelf welke technische records moeten synchroniseren.")
    b.h3("Noodmodus")
    b.p("Een afzonderlijk bevoegde Super Admin kan bij een langdurige ERP-storing een centraal tijdsvenster activeren met reden, begin, einde en audit. Dit laat werken op de laatst gesynchroniseerde agenda toe, maar maakt geen nieuwe contactcenterafspraken en wist geen open opdrachten. Noodmodus heft de cashblokkering niet op.")
    b.h2("4.3 Mijn voorbereiding")
    for item in [
        "Toont de afspraken van de volgende effectieve werkdag, niet simpelweg kalenderdag +1.",
        "Weekends en actieve landfeestdagen worden overgeslagen.",
        "Zichtbaar vanaf een beheerparameter per land; standaard 16:30 in de lokale tijdzone.",
        "Afspraken blijven in bindende ERP/contactcentervolgorde.",
        "De vertegenwoordiger markeert elke voorbereiding expliciet als voorbereid; dit is informatief en blokkeert de latere agenda niet.",
        "Notities zijn zichtbaar voor de vertegenwoordiger en voor management binnen read-only scope.",
        "Aanbevelingen zijn deterministisch en verklaren aankoopdatums, gemiddelde interval, verwachte datum, dagen en hoeveelheid.",
    ]:
        b.bullet(item)
    b.h3("Aanbevelingslogica")
    b.p("De eerste implementatie gebruikt positieve gefactureerde historische regels. Per artikel worden unieke aankoopdatums, gemiddelde interval (of 180 dagen na één aankoop), verwachte herbesteldatum en een horizon van standaard 30 dagen berekend. Feedback relevant/niet relevant verandert de ERP-artikelmaster niet.")
    b.h2("4.4 Mijn agenda en afspraken")
    for item in [
        "Mijn agenda toont alleen afspraken van vandaag en de daarbij toegestane klantgegevens.",
        "De door ERP/contactcenter opgelegde volgorde is bindend; SalesDay optimaliseert de route niet.",
        "De gebruiker kan navigatie op het toestel openen volgens die volgorde.",
        "Een eigen extra afspraak mag alleen voor vandaag worden gemaakt.",
        "Bewerken, dupliceren of annuleren is alleen toegestaan voor de eigen geplande afspraak van vandaag en volgens ERP-redenen.",
        "Een nieuwe of gedupliceerde eigen afspraak komt na de bestaande ERP-volgorde.",
        "Elke contactcenterafspraak krijgt een definitieve uitkomst: uitgevoerd, niet uitgevoerd, verplaatst of geannuleerd volgens de ondersteunde redenflow.",
    ]:
        b.bullet(item)
    b.h3("Niet uitgevoerd of verplaatst")
    b.p("Redenen komen uit het ERP en zijn per land en actiefstatus gecachet. Als een reden uitleg vereist, is de uitleg verplicht. Verplaatst maakt geen toekomstige afspraak en geen impliciete contactcentertaak; ERP/contactcenter blijft eigenaar van toekomstige planning.")
    b.h2("4.5 Klantdossier en klantwijzigingen")
    b.p("SalesDay en Contract delen één centrale BusinessRelation met contacten, adressen en facturatievalidatie. Een vertegenwoordiger ziet normaal volledige details alleen wanneer de klant vandaag op de eigen agenda staat. Voor het toevoegen van een eigen afspraak kan wel in de effectieve klant/teamindex worden gezocht, ook offline.")
    b.h3("Wijzigen bij de klant")
    for item in [
        "De vertegenwoordiger mag tijdens het bezoek officiële klantgegevens rechtstreeks aanpassen.",
        "FieldForce registreert oude en voorgestelde waarde, actor, toestel, afspraak, validatie, tijd en ERP-opdracht.",
        "Een expliciete FieldForce-wijziging heeft prioriteit zolang ze op ERP-bevestiging wacht.",
        "Een concurrerend ERP-event mag de pending wijziging niet stil overschrijven.",
        "Na ERP-bevestiging wordt de ERP-versie opnieuw de effectieve bron.",
    ]:
        b.bullet(item)
    b.h3("BTW, VIES en Peppol")
    for item in [
        "Lokale vorm- en Belgische modulo-97-controle geven onmiddellijk feedback.",
        "Een lokaal geldig nummer mag offline worden opgeslagen.",
        "VIES en Peppol zijn gezaghebbend voor facturatie-identiteit wanneer beschikbaar.",
        "Een conflict tussen officiële bronnen wordt zichtbaar gemaakt en niet stil opgelost.",
        "Onbeschikbaarheid van de externe controle blokkeert verkoop of Contract niet.",
        "Ondertekende documenten bewaren hun oorspronkelijke klant- en facturatiesnapshot.",
    ]:
        b.bullet(item)
    b.h3("Prospect")
    b.p("Wanneer zoeken geen klant oplevert, mag de vertegenwoordiger offline een prospect maken. Bij de eerste Order, Order-Reeds-Geleverd of Factuur wordt de prospect klant. De ERP-opdrachten zijn afhankelijk geordend: eerst relatie/klant, daarna document.")
    b.h2("4.6 Bezoekverslag, lead, opvolging en referentie")
    b.table(
        ["Object", "FieldForce-gedrag", "Eigenaarschap na sync"],
        [
            ["Bezoekverslag", "Eén onveranderlijk verslag bij voltooiing; na dagafsluiting alleen een apart geaudit addendum.", "ERP ontvangt het; FieldForce-zicht blijft beperkt tot originator en managementscope."],
            ["Lead", "Ontstaat in FieldForce en wordt doorgestuurd.", "ERP beheert latere leadstatussen."],
            ["Opvolging", "Ontstaat in FieldForce als concrete vraag.", "ERP/contactcenter plant en beheert latere status."],
            ["Referentie", "Potentiële klant aangebracht door de bezochte klant.", "ERP/contactcenter beslist over omzetting; nooit automatisch FieldForce-prospect."],
        ],
        [1800, 4210, 3350],
    )
    b.h2("4.7 Commerciële documenten")
    b.table(
        ["Type", "Wanneer voorgesteld", "Voorraadeffect"],
        [
            ["Order", "Er is niet genoeg of geen vertegenwoordigersvoorraad.", "Geen onmiddellijke aftrek; levering komt later uit ERP."],
            ["Order-Reeds-Geleverd", "Goederen zijn geleverd, maar ter plaatse factureren mag/kan niet.", "Onmiddellijke aftrek uit vertegenwoordigersvoorraad."],
            ["Factuur", "Voorraad volstaat en ter plaatse factureren is toegestaan.", "Onmiddellijke aftrek uit vertegenwoordigersvoorraad."],
        ],
        [2300, 4010, 3050],
    )
    b.h3("Document maken")
    for item in [
        "Open de afspraak en kies Documenten.",
        "Selecteer artikelen en aantallen uit de officiële ERP-replica; prijs, eenheid, BTW en omschrijving worden gesnapshot.",
        "Controleer het voorgestelde documenttype op basis van voorraad en klantconfiguratie.",
        "Een override is mogelijk, maar vereist een actieve beheerde reden en eventueel verplichte vrije tekst.",
        "Kies de betaalmethode. Alleen ERP-methoden met cash-impact beïnvloeden het kasblad; een gewone Order nooit.",
        "Kies bij carrier-gebonden directe levering de klantdrager.",
        "Laat de klant elk documenttype ondertekenen. Een uitzondering vereist reden en uitleg; de vertegenwoordiger tekent nooit in naam van de klant.",
        "Controleer taal. Standaard is de ERP-klanttaal; vóór ondertekening kan een ondersteunde landtaal worden gekozen.",
        "Toon, deel of print de offline kopie. Na sync verstuurt ERP het officiële document en levert de bezorgstatus terug.",
    ]:
        b.numbered(item)
    b.h3("Nummering en onveranderlijk bewijs")
    b.p("ERP reserveert officiële nummerblokken voor offline gebruik. Toekenning, overslaan, annulering, indiening en bevestiging worden gereconcilieerd. Het bewijs omvat documenthash, snapshots, taal/templateversie, klant-/facturatiegegevens, regels, totalen, handtekening of uitzondering, vertegenwoordiger, toestel en tijdstip.")
    b.h3("Prijs en blokkering offline")
    b.p("De officiële prijs die op het moment van de offline verkoop beschikbaar was blijft geldig. Een latere ERP-prijswijziging herprijst niet stil. Een later ontdekte klantblokkering wordt gelogd en getoond, maar wist een geldige FieldForce-verkoop niet.")
    b.h3("Contract openen vanuit een afspraak")
    b.p("SalesDay mag de bestaande Contractmodule openen met server-gevalideerde klant- en afspraakcontext. Productverkoop en Contract blijven twee afzonderlijke transacties en documentstromen die aan dezelfde afspraak kunnen zijn gekoppeld.")
    b.h2("4.8 Shared Inventory")
    b.p("Inventory is een gedeeld domein voor SalesDay en later Service. Centrale magazijnvoorraad blijft ERP-eigendom. Transit, vertegenwoordigers-/voertuigvoorraad en offline bewegingen worden lokaal beheerd tot ERP-bevestiging.")
    b.h3("Bevoorrading ontvangen")
    for item in [
        "Open de eigen bevoorrading en controleer de verwachte regels.",
        "Registreer per regel de werkelijk ontvangen en beschadigde hoeveelheid; gedeeltelijke ontvangst is toegestaan.",
        "Voeg verplichte digitale bevestiging, handtekening van de vertegenwoordiger en minstens één foto toe.",
        "Tekort, teveel of schade wordt als afzonderlijk verschil naar ERP/backoffice gestuurd.",
        "Beschadigde goederen gaan naar niet-verkoopbare quarantainevoorraad.",
        "Een retry gebruikt dezelfde idempotente ontvangstidentiteit en mag geen dubbele beweging maken.",
    ]:
        b.numbered(item)
    b.h3("Verbruiksgoederen")
    b.p("FieldForce stuurt een aanvraag naar ERP. Na indienen kan de vertegenwoordiger de aanvraag in FieldForce niet goedkeuren, wijzigen of annuleren; magazijn/backoffice beheert picking en latere status.")
    b.h3("Klantlocaties en dragers")
    for item in [
        "Een drager is de fysieke plaats of het object bij de klant waar materiaal staat: kast, rek of bijvoorbeeld een EHBO-koffer.",
        "Een drager kan zelf naar een ERP-artikel verwijzen.",
        "Locaties, sublocaties en dragers kunnen binnen afspraaktoegang worden gemaakt, gewijzigd en gearchiveerd.",
        "Archiveren gebruikt een beheerreden en bewaart alle historiek; fysiek verwijderen is niet toegestaan wanneer historie bestaat.",
        "Een fysieke telling is optioneel. Een verschil maakt onmiddellijk een onveranderlijke correctiebeweging met verplichte reden.",
        "Lot en vervaldatum worden alleen bijgehouden als het ERP-artikel dit vereist. Standaard waarschuwing: 180 dagen.",
    ]:
        b.bullet(item)
    b.h3("Expliciet uitgesloten")
    for item in [
        "Geen overdracht tussen vertegenwoordigers.",
        "Geen retourproces van klant naar vertegenwoordiger binnen SalesDay.",
        "Geen persoonlijke/voertuigvoorraadtelling of correctie door de vertegenwoordiger.",
        "Geen FieldForce-goedkeuring van verbruiksgoederen na indiening.",
    ]:
        b.bullet(item)
    b.h2("4.9 Kasblad en wekelijkse nulregel")
    for item in [
        "Betaalmethoden komen uit ERP en kunnen per land of klant verschillen.",
        "Alleen als cash gemarkeerde methoden verhogen het FieldForce-kassaldo.",
        "Op de eerste effectieve werkdag van de week moet het bevestigde saldo exact nul zijn vóór agenda en voorbereiding openen.",
        "De berekening houdt rekening met landtijdzone, feestdagen en planning; meestal is dit maandag.",
        "Tijdens de blokkering blijven reden/dashboard, kasblad, stortingsproces, synchronisatie en support beschikbaar.",
        "Er bestaat geen Admin- of Super Admin-knop om de kas handmatig vrij te geven.",
        "Vrijgave gebeurt automatisch wanneer ERP/backoffice exact nul bevestigt.",
    ]:
        b.bullet(item)
    b.h2("4.10 Dagafsluiting")
    for item in [
        "Controleer dat elke afspraak een definitieve uitkomst heeft.",
        "Voor niet-uitgevoerd of verplaatst: actieve ERP-reden en verplichte uitleg invullen.",
        "Voltooide afspraken hebben een onveranderlijk bezoekverslag.",
        "Start Dagafsluiting. De lokale sluiting is idempotent en mag bij pending opdrachten doorgaan.",
        "FieldForce meldt dan duidelijk dat synchronisatie nog vereist is.",
        "Na de afspraakdag mag de vertegenwoordiger de afspraak niet meer wijzigen; onopgeloste afspraken worden managementuitzonderingen.",
    ]:
        b.numbered(item)
    b.h2("4.11 Offline, toestel en synchronisatie")
    b.screenshot("diagram-salesday-sync.png", "Figuur 13 — Offline opdrachten, ERP-events en reconciliatie.", "Diagram van SalesDay synchronisatie tussen tablet, FieldForce en ERP.")
    b.h3("Persoonlijk toestel")
    for item in [
        "Elke vertegenwoordiger heeft maximaal één actieve persoonlijke device-registratie.",
        "Een ingetrokken device-ID wordt nooit opnieuw geactiveerd; vervanging krijgt een nieuwe identiteit.",
        "Ruwe encryptiesleutels blijven niet-exporteerbaar op het toestel; de server bewaart alleen een fingerprint.",
        "Remote logout/wipe maakt sessie en registratie ongeldig en bewaart auditbewijs.",
        "Bij hervatten na slaap/lock wordt biometrie of pincode gevraagd; er is geen aparte korte inactiviteitstimer die een klantgesprek onderbreekt.",
    ]:
        b.bullet(item)
    b.h3("Synchronisatie-ervaring")
    for item in [
        "Automatisch zodra verbinding beschikbaar is.",
        "Toont laatste succesvolle sync, aantal pending opdrachten, fouten/conflicten en Nu synchroniseren.",
        "De gebruiker selecteert geen technische records.",
        "Opdrachten blijven bewaard tot expliciete ERP-bevestiging; herhaalde fouten escaleren.",
        "Een timeout is onzeker, nooit automatisch succes. FieldForce vraagt status op of reconcilieert.",
    ]:
        b.bullet(item)
    b.h2("4.12 Management en rapportage")
    b.p("Managementtoegang is operationeel alleen-lezen: Verkoopleider voor team, Sales/Country Manager voor toegewezen landen, Group Manager voor expliciete scope, Admin voor toegewezen scope en Super Admin globaal. Zij werken niet op naam van de vertegenwoordiger in afspraken, klantwijzigingen, verkoop, voorraad, kas of dagafsluiting.")
    b.p("SalesDay toont actuele dagindicatoren: afspraken, documenten, kas, voorraadwaarschuwingen, sync, dagafsluiting en pilotflags. Power BI blijft de officiële bron voor historische KPI's en managementrapportering. De eerste release gebruikt hoogstens een beveiligde externe link; geen iframe-embedding.")
    b.h2("4.13 Mock/UAT en productiegrens")
    for item in [
        "Mockdata is volledig fictief, deterministisch, landdekkend voor BE/NL/DE en gebruikt waar nodig .invalid e-mailadressen.",
        "De seed gebruikt bestaande actieve vertegenwoordigers en maakt geen demo-gebruikers.",
        "De runner weigert productie en databasenamen zonder test/uat/dev/demo/mock/sandbox/local.",
        "De huidige MExT_FieldForce-database wordt daarom bewust geweigerd.",
        "Productie mag nooit terugvallen op mockdata wanneer ERP uitvalt.",
    ]:
        b.bullet(item)


def add_contract(builder: ManualBuilder) -> None:
    b = builder
    b.h1("5. Contract — complete gebruikershandleiding")
    b.p("Contractcalculatie is een native FieldForce-module binnen dezelfde Next.js-shell, aanmelding, rollen, datascope, auditlogging en Prisma/MariaDB-database. De Lovable-tool is vertaald naar FieldForce-code zonder Supabase-runtime, iframe, tweede login of apart rollenmodel.")
    b.screenshot("20-contract-overzicht.png", "Figuur 14 — Contract gebruikt een eigen modulezijbalk; de vastgelegde lokale sessie toont de server-side laadstatus.", "Contractmodule met navigatie en laadstatus.")
    b.h2("5.1 Routes en onderdelen")
    b.table(
        ["Route", "Onderdeel", "Gebruik"],
        [
            ["/contract", "Dashboard", "Aantal berekeningen, jaarlijkse waarde, totale kost en recente dossiers."],
            ["/contract/new", "Nieuwe berekening", "Vierstappenwizard met live preview."],
            ["/contract/calculations", "Berekeningen", "Concepten en ondertekende berekeningen, PDF en ondertekenen."],
            ["/contract/customers", "Klanten", "Klantenkaarten en aantal berekeningen."],
            ["/contract/reporting", "Rapportering", "Operationele verdeling 3/5 jaar, waarde en kost."],
            ["/contract/manage", "Beheer", "Catalogusimport, modelversies en Contractbrief; Admin/Super Admin."],
        ],
        [2300, 2350, 4710],
    )
    b.h2("5.2 Rechten en datascope")
    b.table(
        ["Rol", "Contractscope"],
        [
            ["Representative", "Eigen klanten en berekeningen."],
            ["Sales Leader", "Eigen gegevens en teamgegevens."],
            ["Sales Manager / Admin", "Toegewezen landen; fallback naar eigen land volgens bestaande logica."],
            ["Country Manager", "Eigen land."],
            ["Group Manager / Super Admin", "Alle Contractdata binnen het vastgelegde model."],
        ],
        [2800, 6560],
    )
    b.p("Moduletoegang gebruikt menu.contract.enabled en menu.contract.open. Catalogus-, import- en modelbeheer gebruiken afzonderlijke centrale rechten en zijn in de UI beperkt tot Admin en Super Admin. De server herhaalt de scopecontrole voor overzicht, berekening, documentdownload en beheeractie.")
    b.h2("5.3 Berekening maken")
    b.screenshot("diagram-contract-flow.png", "Figuur 15 — Contractflow met vier wizardstappen en bewijsfase.", "Diagram van de Contractcalculatieflow.")
    b.h3("Stap 1 — Klant")
    for item in [
        "Kies een bestaande klant of Nieuwe manuele klant.",
        "Bij een nieuwe klant is bedrijfsnaam verplicht; contact, e-mail, telefoon, adres en stad zijn optioneel volgens het huidige formulier.",
        "Kies de voorkeurstaal: Nederlands, Frans of Duits. Deze taal bepaalt de klantgerichte brief.",
        "Nieuwe Contractklanten worden via de gedeelde BusinessRelation-richting vastgelegd en behouden een compatibiliteitsprojectie voor bestaande Contractlogica.",
    ]:
        b.bullet(item)
    b.h3("Stap 2 — Artikelen")
    for item in [
        "Zoek op artikelnummer of omschrijving.",
        "Voeg één of meer catalogusartikelen toe.",
        "Vul de hoeveelheid in met ondersteuning voor maximaal drie decimalen in de UI.",
        "Verwijder foutieve regels voordat u doorgaat.",
        "Minstens één geldige regel met hoeveelheid groter dan nul is vereist.",
    ]:
        b.bullet(item)
    b.h3("Stap 3 — Looptijd")
    b.p("Kies een looptijd uit de actieve modelversie. De huidige standaardregels ondersteunen drie en vijf jaar. Het gekozen model en de kortingsregel worden opgeslagen als snapshot, zodat een latere catalogus- of modelwijziging de bestaande berekening niet herschrijft.")
    b.h3("Stap 4 — Afronden")
    for item in [
        "Geef de berekening een herkenbare naam.",
        "Controleer de live preview met subtotaal, korting, jaarlijkse prijs en totale kost.",
        "Sla het concept op. De server rekent opnieuw en is de bron van waarheid.",
        "Download daarna eventueel de interne berekenings-PDF of start ondertekening.",
    ]:
        b.bullet(item)
    b.h2("5.4 Rekenregels")
    b.table(
        ["Element", "Formule"],
        [
            ["Regelbedrag", "quantity × unitPrice"],
            ["Subtotaal", "som van alle regelbedragen"],
            ["Jaarprijs 3 jaar", "subtotal × 0,65"],
            ["Jaarprijs 5 jaar", "subtotal × 0,60"],
            ["Regelkost", "quantity × unitCost"],
            ["Totale kost", "som van alle regelkosten"],
        ],
        [3300, 6060],
    )
    b.p("Geldbedragen gebruiken Prisma Decimal en worden op twee decimalen afgerond volgens de huidige centrale engine. JavaScript floating point is niet de zakelijke rekenbron. De berekening bewaart prijzen, kosten, omschrijvingen, klanttaal, looptijd, korting en modelversie.")
    b.h2("5.5 Concept, PDF en ondertekening")
    b.h3("Interne PDF")
    b.p("Een concept kan een interne berekenings-PDF genereren via de bestaande jsPDF-basis. Deze interne weergave kan kostinformatie bevatten en is niet automatisch dezelfde als de klantbrief.")
    b.h3("Klantgerichte Contractbrief")
    for item in [
        "Ondertekening vereist één actieve DOCX-template in de gekozen klanttaal.",
        "Zonder actieve NL-, FR- of DE-template wordt ondertekening geblokkeerd.",
        "De brief toont geen kostprijs en geen marge. TOTALEKOST is alleen een interne templateparameter; margeparameters bestaan bewust niet.",
        "[PRODUCTLIST] moet alleen in een eigen paragraaf staan en wordt een producttabel met artikelnummer, omschrijving en aantal.",
        "[HANDTEKENING] moet eveneens als eigen paragraaf worden gebruikt.",
        "Elke gegenereerde klantbrief krijgt een oplopende documentversie en wordt nooit overschreven.",
        "De templateversie en parametersnapshot blijven gekoppeld aan de berekening.",
        "Download loopt via een beveiligde API en volgt dezelfde Contractscope als de berekening.",
    ]:
        b.bullet(item)
    b.h2("5.6 Berekeningen, klanten en rapportering")
    b.h3("Berekeningen")
    b.p("De lijst toont nummer, naam, klant, looptijd, jaarlijkse prijs, status en documentacties. Een concept kan worden ondertekend; een klantgerichte PDF kan volgens scope worden gedownload. Ondertekende of gegenereerde documenten blijven immutable bewijs.")
    b.h3("Klanten")
    b.p("Klantenkaarten tonen bedrijfsnaam, primair contact en het aantal Contractberekeningen. In de gedeelde gegevensrichting is BusinessRelation de centrale klant/prospect, terwijl ContractCustomer tijdelijk als compatibiliteitsprojectie blijft bestaan.")
    b.h3("Rapportering")
    b.p("De huidige rapportpagina toont operationele aantallen en waarde/kost, inclusief verdeling voor drie- en vijfjaarscontracten. Dit is geen vrijgeleide voor nieuwe marges, rankings of managementformules; zulke definities vereisen afzonderlijke zakelijke goedkeuring.")
    b.h2("5.7 Catalogusimport")
    b.p("De MEXT_ALL_IN_2026_V1-adapter importeert de aangeleverde Excelstructuur zonder VBA uit te voeren. Import is tweefasig: eerst preview, daarna expliciete bevestiging.")
    b.h3("Validaties")
    for item in [
        "Geldige OOXML/ZIP-signatuur.",
        "Werkblad Input, verborgen Template en werkblad Legende aanwezig.",
        "Tabel MExTBE_Item en vereiste kolommen aanwezig.",
        "Unieke artikelnummers en niet-negatieve prijzen.",
        "Alleen ondersteunde Total Amount-formule; onbekende structuur blokkeert import.",
        "SHA-256/modelversie maakt de import idempotent en verhindert onbedoelde duplicatie.",
    ]:
        b.bullet(item)
    b.h3("Preview")
    b.p("De preview toont gevonden, nieuwe, gewijzigde, gedeactiveerde en ongewijzigde artikelen, plus bronversie en checksum. Pas na bevestiging wordt de import toegepast. Bestaande businessdata wordt niet verwijderd.")
    b.h2("5.8 Templatebeheer")
    for item in [
        "Alleen .docx; .docm, macrobestanden en onbekende parameters worden geblokkeerd.",
        "Templates zijn versieerbaar per taal en bewaren checksum, bron, validatie, uploader en activatie.",
        "Per taal is maximaal één actieve versie tegelijk mogelijk.",
        "Oude versies worden niet verwijderd.",
        "Een nieuwe actieve versie beïnvloedt alleen nieuwe documenten; historisch bewijs blijft aan de oude versie gekoppeld.",
    ]:
        b.bullet(item)
    b.h2("5.9 SalesDay- en ERP-grens")
    b.p("SalesDay opent Contract alleen met server-gevalideerde afspraak- en klantcontext. Een latere facturatiecorrectie herschrijft nooit een reeds ondertekende Contractbrief. Een nog niet ondertekend concept wordt niet automatisch opnieuw opgebouwd; de UI mag een waarschuwing tonen. Contract heeft in de huidige fase geen eigen NAV/Odoo-koppeling, managementgoedkeuring of e-mailflow.")


def add_management(builder: ManualBuilder) -> None:
    b = builder
    b.h1("6. Beheer en gedeelde platformprocessen")
    b.h2("6.1 Gebruikers en teams")
    for item in [
        "Gebruik de bestaande User- en Team-records; maak geen module-eigen gebruikers.",
        "Een team mag zonder primaire verkoopleider bestaan. De database bewaart null, geen dummyaccount.",
        "Een gebruiker heeft één primaire landcontext en kan bijkomende landenscope hebben.",
        "Representative-niveau STARTER, SALES_EXECUTIVE, PROFESSIONAL of EXPERT is een profielwaarde, geen aparte rol.",
        "Profiel- en fotoacties volgen dezelfde persoonlijke datascope als andere gebruikersvelden.",
    ]:
        b.bullet(item)
    b.h2("6.2 Rollen, rechten en overrides")
    for item in [
        "RolePermission bevat de runtime-defaults per rol.",
        "UserPermission bevat alleen expliciete afwijkingen; ontbrekend betekent erven.",
        "Een inactieve rol blijft zichtbaar en bestaande gebruikers behouden de rol, maar nieuwe toekenning is geblokkeerd.",
        "Super Admin krijgt alle geregistreerde rechten; gewone businesslocks blijven gelden.",
        "Nieuwe navigatie vereist rol-default, gebruikersoverride, directe routecontrole, API-controle, vertalingen en documentatie.",
    ]:
        b.bullet(item)
    b.h2("6.3 Modules en featureflags")
    b.p("Coaching-functionaliteit gebruikt centrale modulecodes. SalesDay gebruikt daarnaast server-side flags op Global-, Country-, Team- en User-niveau voor SalesDay, Inventory, offline opdrachten en ERP-writes. Flags gelden tegelijk voor menu, pagina, API, bootstrap en achtergrondverwerking. Een publieke browseromgevingvariabele is geen beveiligingslaag.")
    b.h2("6.4 Parameters")
    b.table(
        ["Domein", "Voorbeelden"],
        [
            ["Coaching", "KPI-definities en doelen, kapstok, criteria, tussentijdse evaluatievragen, mail en profielsync."],
            ["SalesDay", "Voorbereidingstijd per land, documentoverride- en handtekeningredenen, Inventory-redenen, vervalwaarschuwing, Power BI-link."],
            ["Contract", "Actieve modelversie, termijnregels, catalogusimport en actieve brief per taal."],
            ["Platform", "Moduleactivatie, rolstatus, rechten en user overrides."],
        ],
        [2400, 6960],
    )
    b.h2("6.5 Audit en historiek")
    b.p("Zakelijke wijzigingen bewaren actor, tijd, oude/nieuwe toestand en context waar relevant. SalesDay voegt toestel, afspraak, ERP-command en idempotentie toe. Immutable bewijs — ondertekende documenten, bezoekverslagen, addenda, foto's en historische score-snapshots — wordt niet aangepast door latere masterdatawijzigingen.")
    b.h2("6.6 Talen en terminologie")
    for item in [
        "Alle zichtbare teksten ondersteunen Nederlands, Frans en Duits.",
        "UTF-8/utf8mb4 bewaart accenten, umlauts en ß end-to-end.",
        "Gebruik vaste bedrijfstermen zoals Begeleiding, Actiepunt, Prestatiecirkel, Vertegenwoordiger en Verkoopleider.",
        "Documenttaal en klanttaal zijn afzonderlijk van de huidige UI-taal wanneer het proces dat vereist.",
    ]:
        b.bullet(item)


def add_technical(builder: ManualBuilder) -> None:
    b = builder
    b.h1("7. Technisch framework en programma-opbouw")
    b.h2("7.1 Technologiestack")
    b.table(
        ["Laag", "Technologie / versie", "Verantwoordelijkheid"],
        [
            ["Webframework", "Next.js 15.5", "App Router, serverroutes, pagina's en productiebundel."],
            ["UI", "React 19.2 + Tailwind CSS 3.4", "Tablet-first componenten, kaarten, formulieren en responsive gedrag."],
            ["Taal", "TypeScript 5.8 strict", "Getypeerde domeincontracten, services en componenten."],
            ["Database", "Prisma 6.9 + MariaDB/MySQL", "Schema, transacties, relaties, Decimal en migraties."],
            ["Authenticatie", "NextAuth 5 beta; Microsoft Entra richting", "Sessies en gedeelde identiteit."],
            ["Documenten", "jsPDF 4.2, svg2pdf", "Coaching- en Contract-PDF-uitvoer."],
            ["Mail", "Nodemailer 7", "Best-effort FieldForce-e-mail met veilige testmodus."],
            ["Iconen", "Lucide React", "Consistente functionele iconografie."],
        ],
        [1900, 2500, 4960],
    )
    b.h2("7.2 Vier architectuurniveaus")
    for item in [
        "Platform: authenticatie, User, Team, Role, Permission, Module, scope, taal, navigatie, audit, sessies en gedeelde Planningpresentatie.",
        "Domein: Coaching, SalesDay, Inventory, Contract, Service, PST, Reporting en Administration.",
        "Functioneel gebied: bijvoorbeeld Begeleidingen, Mijn agenda, Contractberekeningen.",
        "Workflow: statusovergangen, validatie, rechten, bewijs en eigenaarschap van één businessobject.",
    ]:
        b.numbered(item)
    b.h2("7.3 Bronstructuur")
    b.code("app/                 Next.js pagina's en API-routes\ncomponents/          gedeelde en domeinspecifieke UI\nlib/                 pure businesslogica, rechten en clienthelpers\nlib/server/          serverservices, scope, transacties en integraties\nprisma/              schema en voorwaartse migraties\nscripts/             gerichte validatie-, seed- en beheerscripts\ndocs/ai/             zakelijke bron van waarheid en implementatiebeslissingen\ndocs/technical/      technische database- en deploymentdetails")
    b.h2("7.4 Route- en servicepatroon")
    for item in [
        "Pagina bepaalt presentatie en start een beperkte fetch; zij is geen beveiligingsgrens.",
        "API-route valideert sessie, invoer en het vereiste functionele recht.",
        "Serverservice berekent effectieve scope, controleert lifecycle en voert de businessregel uit.",
        "Prisma-transactie bewaart gerelateerde records atomair.",
        "Audit, bewijs en — voor SalesDay — ERP-outboxcommando worden in dezelfde transactie gekoppeld.",
        "Response geeft alleen data binnen scope terug; queryparameters kunnen nooit scope verbreden.",
    ]:
        b.numbered(item)
    b.h2("7.5 Gegevensmodel en bronnen van waarheid")
    b.table(
        ["Domein", "Centrale modellen / richting", "Belangrijke invariant"],
        [
            ["Platform", "User, Team, RolePermission, UserPermission, AppModule", "Geen module-eigen rollen of gebruikers."],
            ["Coaching", "Intervention, Approval, Score, CoachingAction, ActionPoint, HelpRequest", "Eén record en lifecycle; snapshots bewaren historiek."],
            ["Relatie", "BusinessRelation + contact/adres/validatie/externe link", "Centrale klant/prospect voor SalesDay en Contract."],
            ["SalesDay", "SalesAppointment, preparation, documents, reports, closure", "Pending FieldForce-wijziging blijft tot ERP-bevestiging."],
            ["Inventory", "InventoryLocation, Movement, Balance, Receipt, Discrepancy", "Beweging is immutable en idempotent; saldo is afgeleid."],
            ["Cash", "SalesPaymentMethod, SalesCashBalance, SalesCashEntry", "ERP/backoffice bevestigt; geen manuele override."],
            ["Contract", "ContractCalculation, Line, ModelVersion, LetterTemplate, GeneratedDocument", "Prijs-/klant-/model- en document-snapshots zijn onveranderlijk."],
            ["Integratie", "ErpInboxMessage, ErpOutboxCommand, Dependency, Checkpoint, Incident", "Geen stille command/event loss; unieke ids en reconciliatie."],
        ],
        [1750, 3850, 3760],
    )
    b.h2("7.6 ERP-integratiecontract sales-erp.v1")
    b.p("De integratiegrens is providerneutraal. Een concrete BC/NAV- of Odoo-adapter vertaalt native endpoints, authenticatie, statussen, tijdzones, decimalen, paging en fouten naar stabiele FieldForce-contracten. Onbekende mogelijkheden falen gesloten; adapters mogen niet gokken.")
    b.h3("ERP naar FieldForce")
    for item in [
        "Customers/prospects, contacts, addresses en billing validation.",
        "Appointments en bindende sequence.",
        "Articles, prijzen, BTW, eenheden, lot-/expiryflags.",
        "Commercial history, replenishments, cash balances en payment methods.",
        "Outcome reasons, document categories, customer locations en carrier balances.",
    ]:
        b.bullet(item)
    b.h3("FieldForce naar ERP")
    for item in [
        "Customer/prospect upsert en appointment changes/outcomes.",
        "Visit report/addendum, lead, follow-up en reference.",
        "Order, Order-Reeds-Geleverd en Factuur met snapshots en ondertekening.",
        "Location/carrier upsert, carrier count, replenishment receipt en consumables request.",
        "Day close en staged attachments.",
    ]:
        b.bullet(item)
    b.h3("Idempotentie")
    b.p("De idempotency key is een SHA-256-hash over canonieke JSON met commandtype, business key, context en payload. Objectkeys worden gesorteerd; arrayvolgorde blijft behouden. Retry gebruikt dezelfde command-ID en sleutel. Dezelfde sleutel met andere semantiek is een fout. Afhankelijkheden voorkomen dat een document vóór de nieuwe prospect of eigen afspraak naar ERP gaat.")
    b.h2("7.7 Offline- en toestelarchitectuur")
    for item in [
        "Versleutelde IndexedDB/device store voor scopegebonden replica, drafts en opdrachtwachtrij.",
        "Non-exportable CryptoKey op het toestel; server bewaart fingerprint, tokenhash en provisioningbewijs.",
        "Eén actieve registratie per vertegenwoordiger; revocation bewaart historiek.",
        "Autosave met versies voorkomt overschrijven door vertraagde responses.",
        "Volgende werkdag vereist nieuwe agenda en afgewerkte vorige-dagopdrachten.",
        "Backendcontracten blijven bruikbaar voor een latere native Android-app.",
    ]:
        b.bullet(item)
    b.h2("7.8 Beveiliging")
    b.table(
        ["Risico", "Maatregel"],
        [
            ["IDOR / URL-manipulatie", "Elke API en download herberekent recht, scope en recordrelatie."],
            ["Client-only rechten", "UI is alleen presentatie; serverservice is beslissend."],
            ["Offline replay na intrekking", "Replay vraagt opnieuw server-side autorisatie."],
            ["Dubbele financiële/stockactie", "Business key, idempotency key, unieke databaseconstraints en transactionele outbox."],
            ["Gevoelige lock-screeninfo", "Neutrale notificatie zonder klantnaam of bedrag."],
            ["Verloren tablet", "MDM lock/wipe, device revocation, sessie-invalidatie en nieuwe bootstrap."],
            ["Bestandslek", "Privé uploadroot, geauthenticeerde download en scopecontrole."],
            ["Mock in productie", "Runtime en seed falen gesloten; geen fallback."],
        ],
        [2700, 6660],
    )
    b.h2("7.9 Bestanden en documenten")
    for item in [
        "Contactmomentfoto's staan onder FIELD_FORCE_UPLOAD_ROOT/contact-moments en worden privé gelezen.",
        "Gebruikersfoto's staan onder FIELD_FORCE_UPLOAD_ROOT/user-avatars en kunnen server-side uit Microsoft Graph synchroniseren.",
        "Contractdocumenten en templates bewaren checksum, versie en immutable koppeling.",
        "SalesDay staged attachments krijgen een afspraak- of klantcategorie uit ERP en blijven versleuteld/afgeschermd tot bevestiging.",
        "Back-up en restore moeten database én private uploadroot omvatten.",
    ]:
        b.bullet(item)
    b.h2("7.10 Vertaling en encoding")
    b.p("NL, FR en DE gebruiken één vertaalsysteem. User-facing tekst hoort niet hardcoded in componenten. Database en bestandspaden gebruiken UTF-8/utf8mb4. Reparatie van beschadigde tekst gebeurt alleen na diagnose en met een betrouwbare bron voor de bedoelde tekens.")
    b.h2("7.11 Testarchitectuur")
    b.p("De repository bevat gerichte scripts voor workflow, scope, menu, data access, PDF, mail, Outlook, Contract, ERP-ledger, device-security, offline queue, SalesDay-dagpoort, klanten, afspraken, voorbereiding, commerciële documenten, Inventory, cash en UAT seed. Typecheck, lint, Prisma-validatie en een productiebuild vormen de brede releasecontrole.")


def add_operations(builder: ManualBuilder) -> None:
    b = builder
    b.h1("8. Operaties, test, release en herstel")
    b.h2("8.1 Omgevingen")
    b.table(
        ["Omgeving", "Toegestaan", "Niet toegestaan"],
        [
            ["Development", "Deterministische mockadapter, fictieve demo-seed, gerichte tests.", "Productiegegevens als stilzwijgende fixture."],
            ["Test/UAT", "Fictieve mock of geïsoleerde ERP-testtenant, alle landen/talen.", "Ambigue database of ongecontroleerde productieverbinding."],
            ["Production", "Alleen geaccepteerde echte ERP-adapter, gecontroleerde migraties en featureflags.", "Mockprovider, mockseed of fallback naar demo-inhoud."],
        ],
        [1900, 3730, 3730],
    )
    b.h2("8.2 Migratiepatroon")
    for item in [
        "Preflight: backup/restorebewijs, counts, duplicaten, orphans en capaciteit.",
        "Expand: alleen additieve tabellen, nullable bridges, indexes en flags; oude code blijft werken.",
        "Backfill: begrensde idempotente batches met voortgang en checksums.",
        "Verify: pariteit, Contractregressies, scope, unieke externe ids en fixture rebuild.",
        "Switch: nieuwe canonical reads/writes per gecontroleerde scope.",
        "Observe: incidenten en rollback via featureflag naar compatibiliteitspad.",
        "Contract: oude duplicaatvelden pas in een latere release verwijderen.",
    ]:
        b.numbered(item)
    b.h2("8.3 Minimale releasechecks")
    for item in [
        "TypeScript typecheck en lint zonder nieuwe waarschuwingen.",
        "Prisma validate/generate en migratiestatus.",
        "Menu-, rol-, scope-, API-auth- en direct-route-tests.",
        "Coaching lifecycle, PDF, actiepunt, contactmoment en hulpaanvraag regressies.",
        "Contractberekening, import en brief regressies.",
        "SalesDay offline fault injection, idempotentie, documentnummering, handtekening, Inventory en cash.",
        "NL/FR/DE-pariteit en tablet/PWA-toegankelijkheid.",
        "Echte ERP-testtenant round trip per resource en command vóór productie.",
        "Backup/restore, device-loss, noodmodus en rollback daadwerkelijk oefenen.",
    ]:
        b.bullet(item)
    b.h2("8.4 UAT-framework")
    b.p("Elk land levert minstens één vertegenwoordiger, Verkoopleider, backoffice/magazijngebruiker en Admin/Super Admin. Dezelfde scenario's worden in NL/FR/DE uitgevoerd waar van toepassing.")
    scenarios = [
        "Aanmelden, modulewissel, taal en scopecontrole.",
        "Coaching plannen met/zonder voorafmelding, uitvoeren, autosave, reflectie, akkoord en PDF.",
        "Contactmoment met foto's en delen; Hulpaanvraag met concrete opvolging.",
        "SalesDay bootstrap, voorbereiding na tijdvenster, bindende agenda en eigen afspraak.",
        "Offline klantwijziging, BTW-validatie, prospect en eerste verkoop.",
        "Order, Order-Reeds-Geleverd en Factuur, override, handtekening/uitzondering en print/share.",
        "Replenishment gedeeltelijk/beschadigd, drager, telling, expiry en consumables request.",
        "Cash nulpoort, dag -1 syncpoort, noodmodus en device loss.",
        "Contract bestaande/nieuwe klant, catalogusregels, 3/5 jaar, concept, template en ondertekening.",
        "Management read-only pogingen en directe API/URL buiten scope.",
    ]
    for item in scenarios:
        b.numbered(item)
    b.h2("8.5 Reconciliatieprocedure")
    for item in [
        "Identificeer command, event of entiteit en bewaar alle ledger- en bewijsrecords.",
        "Vergelijk FieldForce command-ID/business key met ERP-acknowledgement of bronidentiteit.",
        "Is ERP-acceptatie bewezen, reconcilieer het bestaande command; maak geen nieuw command met nieuwe identiteit.",
        "Bij ERP-rejectie: bewaar terminale status en toon de operationele fout.",
        "Bij onzekerheid blijft het command retryable of incident open tot ERP uitsluitsel geeft.",
        "Los het incident pas op wanneer bron- en doelbewijs overeenkomen.",
    ]:
        b.numbered(item)
    b.h2("8.6 Backup/restore en toestelverlies")
    b.h3("Backup/restore")
    for item in [
        "Maak een backup met afspraken, documenten, nummers, handtekeningen, stock, cash en pending outbox.",
        "Herstel uitsluitend in non-productie.",
        "Controleer nummers, bewijs, checkpoints, balances en private bestanden.",
        "Voer gerichte tests en productiebuild uit en documenteer operator, tijd en resultaat.",
    ]:
        b.numbered(item)
    b.h3("Verloren toestel")
    for item in [
        "MDM lock/wipe uitvoeren.",
        "DeviceRegistration in FieldForce intrekken.",
        "Bevestigen dat de lokale sleutel niet meer kan synchroniseren.",
        "Vervangtoestel met nieuwe device-identiteit bootstrappen.",
        "Geen onbevestigde command-identiteit hergebruiken zonder gecontroleerde reconciliatie.",
    ]:
        b.numbered(item)


def add_troubleshooting(builder: ManualBuilder) -> None:
    b = builder
    b.h1("9. Probleemoplossing en praktijktips")
    b.table(
        ["Symptoom", "Controle", "Veilige actie"],
        [
            ["Menu ontbreekt", "Module, rolrecht, override, scope en featureflag.", "Laat beheer de effectieve toegang controleren; gebruik geen directe URL als omweg."],
            ["Module niet actief", "Technisch modulebeheer of SalesDay scopeflag.", "Activeer alleen in de bedoelde omgeving en scope."],
            ["Pagina blijft laden", "API-status, sessie, database/migratie en console/serverlog.", "Niet herhaald verversen; verzamel foutbewijs en controleer health/API."],
            ["Begeleiding niet zichtbaar", "Voorafmelding, lifecycle, doelgroep, scope en gedeelde workflowstate.", "Controleer dezelfde record-ID en ververs normale data; maak geen duplicaat."],
            ["Autosave fout", "Netwerk, lokale draftstatus en servervalidatie.", "Laat lokale waarde staan, gebruik retry en sluit pas na flush."],
            ["Akkoordknop geblokkeerd", "Drie reflectievragen betekenisvol ingevuld?", "Sla reflectie op; rapport en acties worden daarna vrijgegeven."],
            ["Actiepunt kan niet sluiten", "actionPointsClose, scope en concrete assignment.", "Laat een bevoegde manager sluiten; wijzig niet de definitie."],
            ["SalesDay dag geblokkeerd", "Open dag -1 commands, cashbalans, device of featureflag.", "Synchroniseer; voor cash wachten op ERP/backoffice nulbevestiging."],
            ["ERP-command blijft open", "Dependencies, acknowledgement, incident en providerstatus.", "Reconcileer dezelfde identiteit; niet opnieuw creëren."],
            ["Contract kan niet ondertekenen", "Actieve taaltemplate, status, klanttaal en scope.", "Activeer een gevalideerde DOCX-template of kies ondersteunde taal."],
            ["Foto ontbreekt", "Privé uploadroot, metadata en scope.", "Controleer storage/backup; behoud gecontroleerde placeholder in rapport."],
        ],
        [2350, 3220, 3790],
    )
    b.h2("9.1 Tips voor vertegenwoordigers")
    for item in [
        "Controleer vóór vertrek de syncstatus en — op de eerste werkdag — het kassaldo.",
        "Laat SalesDay na een offline wijziging zichtbaar aangeven dat synchronisatie nog nodig is.",
        "Gebruik bij klantgegevens de officiële bronnen en leg afwijkingen duidelijk vast.",
        "Laat de klant zelf tekenen; gebruik een uitzondering alleen met de correcte reden en uitleg.",
        "Sluit de dag pas nadat elke afspraak een definitieve uitkomst heeft.",
    ]:
        b.bullet(item)
    b.h2("9.2 Tips voor coaches en managers")
    for item in [
        "Kies voorafmelding bewust; verrassingsbegeleidingen mogen niet uitlekken.",
        "Gebruik een relevante historische referentie, maar laat actuele actiepunten live.",
        "Finaliseren is iets anders dan opslaan of sluiten; controleer verplichte scores en actiepunten.",
        "Lees reflectie read-only en trek alleen de akkoordaanvraag in wanneer een correctie echt nodig is.",
        "Handel Hulpaanvragen af met een concrete uitkomst, niet met een afwijzing of chat.",
    ]:
        b.bullet(item)
    b.h2("9.3 Tips voor beheerders")
    for item in [
        "Activeer modules en SalesDay-flags gefaseerd per scope; documenteer elke pilotwijziging.",
        "Gebruik echte redenconfiguratie per land en archiveer in plaats van verwijderen.",
        "Test alle drie talen na menu- of workflowwijzigingen.",
        "Gebruik nooit de SalesDay mock seed op een ambigu of productieachtig genoemde database.",
        "Bij incidenten: bewijs bewaren, identiteit reconciliëren, geen blinde replay.",
    ]:
        b.bullet(item)


def add_framework(builder: ManualBuilder) -> None:
    b = builder
    b.h1("10. Onderhouds- en wijzigingsframework")
    b.p("Dit hoofdstuk maakt de handleiding herbruikbaar als veranderingskader. Iedere wijziging start bij het zakelijke object en eindigt pas wanneer UI, serverbeveiliging, data, tests en documentatie dezelfde beslissing uitdrukken.")
    b.h2("10.1 Impactanalyse per wijziging")
    b.table(
        ["Vraag", "Te controleren artefacten"],
        [
            ["Wie mag dit zien/doen?", "RolePermission, UserPermission, scopehelpers, route/API, downloads en offline bootstrap."],
            ["Wie bezit de workflow?", "Domeindocument, statushelpers, formulier, service en Planning-link."],
            ["Welke data is canonical?", "Prisma-model, snapshots, compatibiliteitsbridge en externe eigenaar."],
            ["Wat wordt immutable?", "Ondertekening, rapport, addendum, foto, score- of prijsnapshot en audit."],
            ["Werkt het offline?", "Draft, queue, dependency, idempotentie, conflictprioriteit en volgende-dagpoort."],
            ["Welke talen?", "NL/FR/DE sleutels, documenttaal en encoding."],
            ["Hoe testen?", "Gerichte featuretest, scope/IDOR, typecheck, lint en eventueel Prisma/build."],
            ["Welke documentatie bezit de regel?", "Core doc, moduledoc, decision, TODO of history — nooit overal dupliceren."],
        ],
        [3000, 6360],
    )
    b.h2("10.2 Definition of Done")
    for item in [
        "Zakelijke beslissing is expliciet en niet strijdig met het owning document.",
        "Kleinste onderhoudbare bronwijziging gebruikt bestaande componenten en services.",
        "Server-side recht, scope en lifecycle zijn afgedwongen.",
        "Transactie, snapshot, audit en externe command-koppeling zijn waar nodig atomair.",
        "NL/FR/DE en UTF-8 zijn gecontroleerd.",
        "Gerichte tests, typecheck en risicogebonden bredere checks zijn uitgevoerd.",
        "Owning document is bijgewerkt; open werk staat niet verstopt in permanente besluiten.",
        "Geen productie- of externe wijziging zonder expliciete bevoegdheid en releaseproces.",
    ]:
        b.bullet(item)
    b.h2("10.3 Nieuwe module integreren")
    for item in [
        "Definieer domeineigenaarschap, businessobjecten en status van elk functioneel gebied.",
        "Hergebruik User, Team, Role, Permission, Module, taal, shell, audit en scope.",
        "Voeg navigatie toe met rechtencatalogus, defaults, overrides en tests.",
        "Ontwerp serverroutes en services vóór UI-only gedrag.",
        "Maak datamodel additief en bepaal canonical versus replica/snapshot.",
        "Definieer integratieports providerneutraal en laat onbekende adapters gesloten falen.",
        "Bouw featureflags, testfixtures, UAT-scenario's en rollback vóór productieactivatie.",
        "Documenteer gebruikerswerkwijze en technische runbook samen met de bronwijziging.",
    ]:
        b.numbered(item)
    b.h2("10.4 Documentonderhoud")
    b.p("Werk deze handleiding bij wanneer een zichtbare workflow, rol, status, rekenregel, datagrens of productiepoort wijzigt. Vervang screenshots alleen wanneer het scherm materieel verandert. Noteer steeds documentdatum, broncommit of release, actieve omgeving en bekende beperkingen. Laat niet-actieve of externe-gatefuncties duidelijk gemarkeerd staan.")


def add_appendices(builder: ManualBuilder) -> None:
    b = builder
    b.h1("11. Bijlagen")
    b.h2("11.1 Begrippenlijst")
    b.table(
        ["Term", "Betekenis"],
        [
            ["Begeleiding", "Gestructureerde coachinginterventie met voorbereiding, scores, actiepunten en goedkeuring."],
            ["Gecoachte persoon", "De medewerker die de Begeleiding ontvangt en de goedkeuring/reflection uitvoert."],
            ["Scope", "De combinatie van eigen gebruiker, team, land(en) of globale organisatorische toegang."],
            ["Snapshot", "Onveranderlijke kopie van relevante gegevens op het zakelijke moment."],
            ["BusinessRelation", "Centrale klant/prospectidentiteit voor SalesDay en Contract."],
            ["Carrier / drager", "Fysieke plaats of object bij de klant waarin materiaal wordt bijgehouden."],
            ["Outbox", "Duurzame wachtrij van FieldForce-opdrachten richting ERP."],
            ["Inbox", "Duurzame, gededupliceerde verwerking van ERP-events."],
            ["Idempotentie", "Dezelfde businessopdracht opnieuw uitvoeren zonder een tweede zakelijke uitkomst."],
            ["Reconciliatie", "Bewijs tussen FieldForce en ERP vergelijken en één bestaande identiteit afronden."],
            ["Dag -1", "Eerdere businessdatum met nog niet door ERP bevestigde opdrachten."],
            ["Effectieve werkdag", "Werkdag bepaald met tijdzone, weekend, feestdag en planning."],
        ],
        [2500, 6860],
    )
    b.h2("11.2 Belangrijkste bronregister")
    sources = [
        "AGENTS.md — repositorybrede ontwikkel- en veiligheidsregels",
        "docs/ai/00_PROJECT.md — missie, gebruikers, landen, talen en ERP-richting",
        "docs/ai/01_ARCHITECTURE.md — platform-, domein- en workflowarchitectuur",
        "docs/ai/02_DATABASE.md — gegevensprincipes, modellen en migratiestatus",
        "docs/ai/03_ROLES.md — rollen, scope en effectieve rechten",
        "docs/ai/04_UI_GUIDELINES.md — tablet-first FieldForce UI",
        "docs/ai/modules/Coaching/* — Coachingbeslissingen en workflows",
        "docs/ai/modules/Salesday/DECISIONS.md — goedgekeurd SalesDay-gedrag",
        "docs/ai/modules/Salesday/IMPLEMENTATION-PLAN.md — milestones en productiegates",
        "docs/ai/modules/Salesday/ERP-CONTRACT.md — sales-erp.v1",
        "docs/contract-calculation.md — Contractberekening, import en brief",
        "prisma/schema.prisma — persistent datamodel",
        "lib/app-switcher.ts — domeinen, navigatie en menu-rechten",
        "components/salesday/salesday-workspace.tsx — huidige SalesDay-presentatie",
        "components/contract/contract-workspace.tsx — huidige Contract-presentatie",
    ]
    for item in sources:
        b.bullet(item)
    b.h2("11.3 Open externe beslissingen en beperkingen")
    for item in [
        "Exacte BC/NAV API, middleware, authenticatie, event- en acknowledgementmechaniek.",
        "Odoo-adapter, timing en doelinterface.",
        "VIES/Peppol-eigenaarschap en credentials.",
        "ERP-reservering van officiële documentnummers en bezorgbevestiging.",
        "Finale Android-tablet- en printerkeuze; Epson WF-110 is alleen referentie.",
        "Juridische/Finance/DPO-retentieperioden.",
        "Power BI embedding, licentie, SSO en row-level security; beveiligde link volstaat initieel.",
        "Volledige Group Manager- en Service Operator-Coachingdefinitie.",
        "Retraining, Salestraining en Coaching-rapportage.",
        "Volledige tussentijdse evaluatie-akkoordflow en definitieve PDF.",
    ]:
        b.bullet(item)
    b.h2("11.4 Productiecheck SalesDay")
    checks = [
        "Echte ERP round trips voor elk resource- en commandtype bewezen.",
        "Geen open P0/P1 scope-, data-, idempotentie-, financiële of stockdefecten.",
        "BE/NL/DE UAT-groepen hebben ondertekend.",
        "Migratierepetitie en rollback doorlopen.",
        "Backup/restore en device-loss/MDM geoefend.",
        "Noodmodus geoefend en rollen/incidentcommunicatie bevestigd.",
        "Geen mockprovider of mockbusinessdata in productie.",
        "Pending commands, nummers, handtekeningen, stock en cash blijven bij rollback intact.",
    ]
    for item in checks:
        b.bullet("☐ " + item)
    b.callout("Eindboodschap", "FieldForce is één digitaal werkplatform. De kwaliteit ervan komt niet alleen uit schermen, maar uit één workflow per object, expliciete gegevensbronnen, server-side scope, onveranderlijk bewijs, betrouwbare synchronisatie en documentatie die dezelfde zakelijke waarheid vertelt.", "success")


def build() -> Path:
    save_diagrams()
    builder = ManualBuilder()
    builder.title_page()
    builder.h1("Inhoud en navigatie")
    builder.p("De definitieve DOCX krijgt na generatie een statische, klikbare inhoudsopgave. De PDF volgt dezelfde hoofdstukstructuur. Gebruik in Word daarnaast het navigatievenster om op kopniveau te springen.")
    builder.toc()
    builder.page_break()
    add_intro(builder)
    add_platform(builder)
    add_coaching(builder)
    add_salesday(builder)
    add_contract(builder)
    add_management(builder)
    add_technical(builder)
    add_operations(builder)
    add_troubleshooting(builder)
    add_framework(builder)
    add_appendices(builder)
    builder.save()
    return OUTPUT_FILE


if __name__ == "__main__":
    output = build()
    print(output)
