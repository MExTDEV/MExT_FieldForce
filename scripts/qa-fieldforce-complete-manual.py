from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image, ImageDraw, ImageFont
import pypdfium2 as pdfium
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "docx" / "MExT_FieldForce_Complete_Handleiding.docx"
PDF = ROOT / "output" / "pdf" / "MExT_FieldForce_Complete_Handleiding.pdf"
QA_DIR = ROOT / "tmp" / "manual-qa"
PAGES_DIR = QA_DIR / "pages"
CONTACT_DIR = QA_DIR / "contact-sheets"
REPORT = QA_DIR / "report.json"


def font(size: int, bold: bool = False):
    path = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()


def render_pdf() -> list[Path]:
    PAGES_DIR.mkdir(parents=True, exist_ok=True)
    document = pdfium.PdfDocument(str(PDF))
    outputs: list[Path] = []
    for index in range(len(document)):
        page = document[index]
        bitmap = page.render(scale=1.55, rotation=0)
        image = bitmap.to_pil().convert("RGB")
        output = PAGES_DIR / f"page-{index + 1:03d}.png"
        image.save(output, optimize=True)
        outputs.append(output)
        page.close()
    document.close()
    return outputs


def image_metrics(paths: list[Path]) -> tuple[list[dict], list[int], list[int]]:
    metrics: list[dict] = []
    sparse: list[int] = []
    edge_flags: list[int] = []
    for number, path in enumerate(paths, 1):
        with Image.open(path).convert("L") as image:
            thumb = image.resize((max(1, image.width // 8), max(1, image.height // 8)))
            pixels = list(thumb.get_flattened_data())
            ink = sum(1 for value in pixels if value < 245) / len(pixels)
            border = list(image.crop((0, 0, image.width, 4)).get_flattened_data())
            border += list(image.crop((0, image.height - 4, image.width, image.height)).get_flattened_data())
            border += list(image.crop((0, 0, 4, image.height)).get_flattened_data())
            border += list(image.crop((image.width - 4, 0, image.width, image.height)).get_flattened_data())
            edge_ink = sum(1 for value in border if value < 230) / len(border)
        metrics.append({"page": number, "ink_ratio": round(ink, 5), "edge_ink_ratio": round(edge_ink, 5)})
        if ink < 0.006:
            sparse.append(number)
        if edge_ink > 0.003:
            edge_flags.append(number)
    return metrics, sparse, edge_flags


def contact_sheets(paths: list[Path]) -> list[Path]:
    CONTACT_DIR.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []
    per_sheet = 12
    columns = 3
    cell_w, cell_h = 330, 445
    for start in range(0, len(paths), per_sheet):
        batch = paths[start:start + per_sheet]
        rows = (len(batch) + columns - 1) // columns
        sheet = Image.new("RGB", (columns * cell_w, rows * cell_h + 42), "#E8EEF5")
        draw = ImageDraw.Draw(sheet)
        title = f"MExT FieldForce – pagina's {start + 1}–{start + len(batch)}"
        draw.text((14, 8), title, font=font(20, True), fill="#0B315B")
        for offset, path in enumerate(batch):
            with Image.open(path).convert("RGB") as page:
                page.thumbnail((300, 390), Image.Resampling.LANCZOS)
                col = offset % columns
                row = offset // columns
                x = col * cell_w + (cell_w - page.width) // 2
                y = 42 + row * cell_h + 24
                sheet.paste(page, (x, y))
                draw.rectangle((x - 1, y - 1, x + page.width, y + page.height), outline="#8FA4BA", width=1)
                draw.text((col * cell_w + 14, 42 + row * cell_h + 2), f"Pagina {start + offset + 1}", font=font(15, True), fill="#172033")
        output = CONTACT_DIR / f"contact-{start + 1:03d}-{start + len(batch):03d}.png"
        sheet.save(output, optimize=True)
        outputs.append(output)
    return outputs


def audit_pdf() -> dict:
    reader = PdfReader(str(PDF))
    texts = [(page.extract_text() or "") for page in reader.pages]
    empty = [index + 1 for index, text in enumerate(texts) if len(text.strip()) < 80]
    joined = "\n".join(texts)
    try:
        outline_count = len(reader.outline)
    except Exception:
        outline_count = None
    return {
        "pages": len(reader.pages),
        "characters": len(joined),
        "short_text_pages": empty,
        "outline_root_entries": outline_count,
        "metadata": {str(key): str(value) for key, value in (reader.metadata or {}).items()},
        "mojibake_markers": {marker: joined.count(marker) for marker in ["Ã", "Â", "â€"]},
    }


def audit_docx() -> dict:
    with ZipFile(DOCX) as archive:
        corrupt_member = archive.testzip()
        document_xml = archive.read("word/document.xml").decode("utf-8")
    document = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    headings = [p for p in document.paragraphs if p.style and p.style.name.startswith("Heading")]
    alt_texts: list[str] = []
    for shape in document.inline_shapes:
        descr = shape._inline.docPr.get("descr")
        if descr:
            alt_texts.append(descr)
    return {
        "zip_corrupt_member": corrupt_member,
        "paragraphs": len(document.paragraphs),
        "headings": len(headings),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "images_with_alt_text": len(alt_texts),
        "characters": len(text),
        "toc_field_present": "TOC \\o" in document_xml,
        "update_fields_enabled": "updateFields" in ZipFile(DOCX).read("word/settings.xml").decode("utf-8"),
        "mojibake_markers": {marker: text.count(marker) for marker in ["Ã", "Â", "â€"]},
    }


def main() -> None:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    pages = render_pdf()
    metrics, sparse, edge_flags = image_metrics(pages)
    contacts = contact_sheets(pages)
    report = {
        "docx": audit_docx(),
        "pdf": audit_pdf(),
        "visual": {
            "rendered_pages": len(pages),
            "sparse_pages": sparse,
            "edge_flags": edge_flags,
            "page_metrics": metrics,
            "contact_sheets": [str(path.relative_to(ROOT)) for path in contacts],
        },
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(REPORT)
    print(json.dumps(report, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
